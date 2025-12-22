import os
import io
import tempfile
from PIL import Image
from docx import Document
import comtypes.client
from PySide6.QtWidgets import QLineEdit, QTextEdit, QLabel
from controllers.ReporteController import ReporteController
from controllers.EmpresaController import EmpresaController
from utils.common.metodosGenerales import MetodosGenerales
from utils.common.rutasarchivos import resource_path
from PySide6.QtCore import Qt
import psutil
import shutil
import time

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.shared import RGBColor

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class DocxReport:
    def __init__(self):
        self.document = Document()
        self._setup_document()

    def _setup_document(self):
        # Configurar el tamaño de la página a A4
        sections = self.document.sections
        for section in sections:
            section.page_height = Mm(297)  # Altura A4 en milímetros
            section.page_width = Mm(210)   # Ancho A4 en milímetros
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.different_first_page_header_footer = True  # Primera página diferente

        # Configurar el estilo de fuente predeterminado a Arial
        styles = self.document.styles
        styles['Normal'].font.name = 'Arial'
        styles['Normal'].font.size = Pt(12)

        # Configurar el estilo de los encabezados para usar Arial y color negro
        for heading_style in ['Heading 1', 'Heading 2']:
            style = styles[heading_style]
            style.font.name = 'Arial'
            style.font.color.rgb = RGBColor(0, 0, 0)  # Color negro

    def add_first_page(self, img_blob, titulo, lugar, fecha):
        # Añadir párrafos vacíos para intentar centrar verticalmente
        for _ in range(8):  # Ajusta este número según sea necesario
            self.document.add_paragraph()

        # Añadir la imagen centrada
        if img_blob:
            try:
                img = Image.open(io.BytesIO(img_blob))
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='PNG')
                img_byte_arr.seek(0)

                paragraph = self.document.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(img_byte_arr, width=Inches(2.0))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Error al cargar la imagen: {e}")

        # Añadir título centrado
        title_paragraph = self.document.add_paragraph()
        title_run = title_paragraph.add_run(titulo)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.color.rgb = RGBColor(0, 0, 0)  # Color negro
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Añadir lugar centrado
        lugar_paragraph = self.document.add_paragraph(lugar)
        lugar_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Añadir fecha centrada
        fecha_paragraph = self.document.add_paragraph(f"MES: {fecha}")
        fecha_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_second_page(self, data, datos_firma):
        # Añadir una nueva página para la información adicional
        self.document.add_page_break()

        # Añadir título centrado
        title_paragraph = self.document.add_paragraph(data[3])
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Añadir información adicional con tabulaciones
        additional_info = [
            f"Para\t\t\t: {data[6]}",
            f"De\t\t\t: {data[7]}",
            f"Cc\t\t\t: {data[8]}",
            f"Fecha\t\t\t: {data[5]}",
            f"Asunto\t\t: {data[9]}"
        ]

        for info in additional_info:
            paragraph = self.document.add_paragraph(info)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        self.document.add_paragraph()
        # Añadir el contenido del cuerpo del capítulo
        body_paragraph = self.document.add_paragraph(data[10])
        body_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justificar el texto
        for _ in range(10):  # Ajusta este número según sea necesario
            self.document.add_paragraph()
        # Añadir firma en la esquina inferior izquierda
        if datos_firma and datos_firma[6]:
            try:
                img = Image.open(io.BytesIO(datos_firma[6]))
                img_path = tempfile.mktemp(suffix='.png')
                img.save(img_path)

                # Añadir la imagen de la firma
                paragraph = self.document.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(img_path, width=Inches(1.0))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Eliminar el archivo temporal
                os.remove(img_path)
            except Exception as e:
                print(f"Error al cargar la imagen de la firma: {e}")

        # Añadir el texto de la firma
        if datos_firma:
            signature_text = f"{datos_firma[2]}"
            signature_paragraph = self.document.add_paragraph(signature_text)
            signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_third_page(self, title, img_blob):
        self.document.add_page_break()

        # Añadir título centrado y en negrita
        heading = self.document.add_heading(level=1)
        run = heading.add_run(title)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Color negro
        run.font.size = Pt(20)
        heading.alignment = 1  # Centrar el título

        # Añadir imagen si existe
        if img_blob:
            try:
                img = Image.open(io.BytesIO(img_blob))
                temp_img_path = tempfile.mktemp(suffix='.png')
                img.save(temp_img_path, format='PNG')

                # Obtener el ancho y alto de la página en pulgadas
                section = self.document.sections[0]
                page_width = section.page_width.inches
                page_height = section.page_height.inches
                margins = (section.left_margin.inches, section.right_margin.inches,
                        section.top_margin.inches, section.bottom_margin.inches)
                available_width = page_width - margins[0] - margins[1]
                available_height = page_height - margins[2] - margins[3]

                # Calcular el tamaño de la imagen para ajustarse al espacio disponible
                img_width, img_height = img.size
                width_ratio = available_width / img_width
                height_ratio = available_height / img_height
                scale = min(width_ratio, height_ratio)

                new_width = img_width * scale
                new_height = img_height * scale

                # Añadir la imagen al documento con el nuevo tamaño
                self.document.add_picture(temp_img_path, width=Inches(new_width), height=Inches(new_height))

                # Centrar la imagen
                paragraph = self.document.paragraphs[-1]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Eliminar el archivo temporal
                os.remove(temp_img_path)
            except Exception as e:
                paragraph = self.document.add_paragraph("Error al cargar la imagen")
                paragraph.alignment = 1  # Centrar el texto
        else:
            paragraph = self.document.add_paragraph("Error al cargar la imagen")
            paragraph.alignment = 1  # Centrar el texto

    def add_instrumentation_page(self, lista_prismas, lista_Inclinometros, lista_piezometros, lista_celdas, lista_acelerografos, lista_sondajes):
        # Añadir una nueva página
        self.document.add_page_break()

        # Añadir título de la sección
        heading = self.document.add_heading('1. Instrumentación Geotécnica', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Lista de instrumentos con sus respectivos datos y títulos
        instrumentos = [
            (lista_prismas, 'Prismas'),
            (lista_Inclinometros, 'Inclinómetros'),
            (lista_piezometros, 'Piezómetros'),
            (lista_celdas, 'Celdas'),
            (lista_acelerografos, 'Acelerógrafos'),
            (lista_sondajes, 'Sondajes TDR')
        ]

        # Contador para la numeración de las secciones
        section_counter = 1
        for datos, titulo in instrumentos:
            if datos:  # Solo generar la tabla si hay datos
                # Añadir subtítulo
                subheading = self.document.add_heading(f'1.{section_counter} {titulo}', level=2)
                subheading.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Crear tabla
                table = self.document.add_table(rows=1, cols=5)
                table.style = 'Table Grid'

                # Añadir encabezados con formato en negrita y fondo gris
                headers = ['ID', 'Este', 'Norte', 'Elevación', 'Ubicación']
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    header_cells[i].paragraphs[0].runs[0].font.bold = True
                    # Establecer el color de fondo de la celda a gris claro
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'D3D3D3')  # Color gris claro
                    header_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

                    # Centrar verticalmente la celda
                    tc = header_cells[i]._tc
                    tcPr = tc.get_or_add_tcPr()
                    vAlign = OxmlElement('w:vAlign')
                    vAlign.set(qn('w:val'), 'center')
                    tcPr.append(vAlign)

                # Añadir datos
                for id_instrumento, este, norte, elevacion, ubicacion in datos:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(id_instrumento)
                    row_cells[1].text = str(este)
                    row_cells[2].text = str(norte)
                    row_cells[3].text = str(elevacion)
                    row_cells[4].text = ubicacion
                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        # Centrar verticalmente la celda
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        vAlign = OxmlElement('w:vAlign')
                        vAlign.set(qn('w:val'), 'center')
                        tcPr.append(vAlign)

                # Incrementar el contador de secciones
                section_counter += 1

    def add_interpretation_page(self, listaimagenes):
        # Añadir una nueva página
        self.document.add_page_break()

        # Añadir título de la sección
        heading = self.document.add_heading('2. Interpretación', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for imagen in listaimagenes:
            title = imagen[5]
            img_blob = imagen[4]
            description = imagen[6]

            # Añadir título de la imagen
            title_paragraph = self.document.add_paragraph(title)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Añadir imagen si existe
            if img_blob:
                try:
                    img = Image.open(io.BytesIO(img_blob))
                    if img.mode in ('RGBA', 'P'):
                        img = img.convert('RGB')

                    temp_img_path = tempfile.mktemp(suffix='.jpg')
                    img.save(temp_img_path, format='JPEG', quality=85, optimize=True)

                    # Añadir la imagen al documento centrada
                    paragraph = self.document.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(temp_img_path, width=Inches(5.0))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Eliminar el archivo temporal
                    os.remove(temp_img_path)
                except Exception as e:
                    error_paragraph = self.document.add_paragraph(f'Error al cargar la imagen: {e}')
                    error_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Añadir descripción
            description_paragraph = self.document.add_paragraph(description)
            description_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            self.document.add_paragraph()  # Añadir espacio entre imágenes

    def add_conclusions_page(self, conclusions, recommendations):
        # Añadir una nueva página
        self.document.add_page_break()

        # Añadir título de la sección
        heading = self.document.add_heading('3. Conclusiones y Recomendaciones', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Añadir conclusiones
        conclusions_heading = self.document.add_heading('Conclusiones', level=2)
        conclusions_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        conclusion_paragraphs = conclusions.split('\n')
        for paragraph in conclusion_paragraphs:
            self.document.add_paragraph(paragraph)
            self.document.add_paragraph()  # Añadir espacio entre párrafos

        # Añadir recomendaciones
        recommendations_heading = self.document.add_heading('Recomendaciones', level=2)
        recommendations_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        recommendation_paragraphs = recommendations.split('\n')
        for paragraph in recommendation_paragraphs:
            self.document.add_paragraph(paragraph)
            self.document.add_paragraph()  # Añadir espacio entre párrafos
    
    def add_header(self, img_blob, header_text):
        # Obtener la sección actual
        section = self.document.sections[-1]

        # Crear un encabezado
        header = section.header

        # Añadir una tabla de una fila y dos columnas para el encabezado
        table = header.add_table(rows=1, cols=2, width=Inches(6.0))  # Especificar el ancho de la tabla

        # Añadir imagen a la izquierda
        if img_blob:
            try:
                img = Image.open(io.BytesIO(img_blob))
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='PNG')
                img_byte_arr.seek(0)

                cell = table.cell(0, 0)
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(img_byte_arr, width=Inches(1.0))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Centrar verticalmente la celda
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
            except Exception as e:
                print(f"Error al cargar la imagen: {e}")

        # Añadir texto a la derecha
        cell = table.cell(0, 1)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header_text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Centrar verticalmente la celda
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)

    def add_footer(self, footer_text):
        # Obtener la sección actual
        section = self.document.sections[-1]

        # Crear un pie de página
        footer = section.footer

        # Añadir una tabla de una fila y dos columnas para el pie de página
        table = footer.add_table(rows=1, cols=2, width=Inches(6.0))  # Especificar el ancho de la tabla

        # Añadir texto a la izquierda
        cell = table.cell(0, 0)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(footer_text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Centrar verticalmente la celda
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)

        # Añadir número de página a la derecha
        cell = table.cell(0, 1)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        # Añadir un campo de número de página
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        run._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

        run.font.name = 'Arial'
        run.font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Centrar verticalmente la celda
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)

    def save_docx(self, file_path):
        self.document.save(file_path)

    
    def convert_to_pdf(self, docx_path, pdf_path):
        word = None
        doc = None
        intentos = 3
        
        for intento in range(intentos):
            try:
                import win32com.client
                # Limpiar procesos residuales
                self._limpiar_procesos_word()
                # Crear instancia INDEPENDIENTE de Word
                word = win32com.client.DispatchEx("Word.Application")
                word.Visible = False
                word.DisplayAlerts = False

                # Abrir documento
                doc = word.Documents.Open(docx_path)
                
                # Actualizar campos e índices
                doc.TablesOfContents(1).Update() if doc.TablesOfContents.Count > 0 else None
                doc.Fields.Update()
                
                # Guardar como PDF
                doc.SaveAs(pdf_path, FileFormat=17)
                return True

            except Exception as e:
                if intento == intentos-1:
                    print(f"Error crítico al generar PDF: {str(e)}")
                time.sleep(2)
            finally:
                # Cierre seguro en orden inverso
                if doc:
                    try:
                        doc.Close(SaveChanges=0)
                    except: pass
                if word:
                    try:
                        word.Quit()
                    except: pass
                self._limpiar_procesos_word()

    def _limpiar_procesos_word(self):
        """Elimina solo procesos Word creados por esta instancia"""
        try:
            for proc in psutil.process_iter(['pid', 'name', 'cmdline']):
                try:
                    if proc.info['name'] == 'WINWORD.EXE' and \
                    any(str(os.getpid()) in cmd for cmd in (proc.info['cmdline'] or [])):
                        proc.kill()
                except (psutil.NoSuchProcess, psutil.AccessDenied):
                    continue
        except Exception as e:
            print(f"Error limpiando procesos: {str(e)}")

class ReporteGeneral:
    @staticmethod
    def guardarFormularioReporte(main, idproyecto, file_patch):
        data = ReporteGeneral.obtenerValoresFormulario(main, idproyecto, file_patch)
        ReporteController.ctrlGuardarInformacionReporte(data)

    @staticmethod
    def cargarDataFormulariosReporte(main, idproyecto):
        general = ReporteController.ctrlListarDatosReporteGeneral(idproyecto)
        if general:
            main.findChild(QLineEdit, "input_encabezado").setText(general[2])
            main.findChild(QLineEdit, "input_pie_pagina").setText(general[3])
            main.findChild(QLineEdit, "input_titulo").setText(general[4])
            main.findChild(QLineEdit, "input_lugar").setText(general[5])
            main.findChild(QLineEdit, "input_mes_reporte").setText(general[6])
            main.findChild(QLineEdit, "input_para").setText(general[7])
            main.findChild(QLineEdit, "input_de").setText(general[8])
            main.findChild(QLineEdit, "input_cc").setText(general[9])
            main.findChild(QLineEdit, "input_asunto").setText(general[10])
            main.findChild(QTextEdit, "input_descripcion").setPlainText(general[11])
            main.findChild(QTextEdit, "input_conclusiones").setPlainText(general[12])
            main.findChild(QTextEdit, "input_recomendaciones").setPlainText(general[13])
            if general[14]:
                pixmap = MetodosGenerales.convertir_blob_a_pixmap(general[14])
                scaled_pixmap = pixmap.scaled(120, 120, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                main.findChild(QLabel, "label_imagen").setPixmap(scaled_pixmap)

    @staticmethod
    def obtenerValoresFormulario(main, idproyecto, file_patch):
        encabezado = main.findChild(QLineEdit, "input_encabezado").text()
        pie = main.findChild(QLineEdit, "input_pie_pagina").text()
        titulo = main.findChild(QLineEdit, "input_titulo").text()
        lugar = main.findChild(QLineEdit, "input_lugar").text()
        mes = main.findChild(QLineEdit, "input_mes_reporte").text()
        para = main.findChild(QLineEdit, "input_para").text()
        de = main.findChild(QLineEdit, "input_de").text()
        cc = main.findChild(QLineEdit, "input_cc").text()
        asunto = main.findChild(QLineEdit, "input_asunto").text()
        descripcion = main.findChild(QTextEdit, "input_descripcion").toPlainText()
        conclusiones = main.findChild(QTextEdit, "input_conclusiones").toPlainText()
        recomendaciones = main.findChild(QTextEdit, "input_recomendaciones").toPlainText()
        imagen_componente = file_patch
        data = [idproyecto, encabezado, pie, titulo, lugar, mes, para, de, cc, asunto, descripcion, conclusiones, recomendaciones, imagen_componente]
        return data

    def generarReporte(main, idproyecto, file_patch, id_componente):
        ruta_word = resource_path('modules/reportes/reporte_general.docx')
        ruta_pdf = resource_path('modules/reportes/reporte_general.pdf')
        # Obtener Información de la empresa
        datos_empresa, lic = EmpresaController.ctrlObtenerDatosEmpresa()
        data = ReporteGeneral.obtenerValoresFormulario(main, idproyecto, file_patch)
        datosFirma = ReporteController.ctrlObtenerDatosFirma(idproyecto)
        lista_prismas = ReporteController.ctrlObtenerListaPrismas(idproyecto,id_componente)
        lista_Inclinometros = ReporteController.ctrlObtenerListaInclinometros(idproyecto,id_componente)
        lista_piezometros = ReporteController.ctrlObtenerListaPiezometros(idproyecto,id_componente)
        lista_celdas = ReporteController.ctrlObtenerListaCeldas(idproyecto,id_componente)
        lista_acelerografos = ReporteController.ctrlObtenerListaAcelerografos(idproyecto,id_componente)
        lista_sondajes = ReporteController.ctrlObtenerListaSondajesTDR(idproyecto,id_componente)
        imagenesReporte = ReporteController.ctrlObtenerListaImagenesReporte(id_componente)
        # Crear instancia de DocxReport
        docx_report = DocxReport()
        # Generar el contenido del documento Word
        docx_report.add_first_page(datos_empresa[5], data[3], data[4], data[5])
        # Añadir contenido a partir de la segunda página
        docx_report.add_second_page(data, datosFirma)
        # Añadir contenido a partir de la segunda página
        docx_report.add_third_page(data[4], data[13])
        # Página de Instrumentación Geotécnica
        docx_report.add_instrumentation_page(lista_prismas,lista_Inclinometros,lista_piezometros,lista_celdas,lista_acelerografos,lista_sondajes)
        # Página de Interpretación
        if imagenesReporte:
            docx_report.add_interpretation_page(imagenesReporte)
        docx_report.add_conclusions_page(data[11], data[12])
        # Añadir encabezado y pie de página
        docx_report.add_header(datos_empresa[5], data[1])
        docx_report.add_footer(data[2])
        # Guardar el documento Word
        docx_report.save_docx(ruta_word)
        if not docx_report.convert_to_pdf(ruta_word, ruta_pdf):
            raise Exception("Error al generar PDF después de 3 intentos")
    