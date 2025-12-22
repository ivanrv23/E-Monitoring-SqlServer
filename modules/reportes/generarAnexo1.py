import os
import shutil
import time
import pytz
from datetime import datetime
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from io import BytesIO
from PIL import Image
from PySide6.QtWidgets import (QMessageBox, QLineEdit, QPlainTextEdit, QLabel, QRadioButton, QComboBox, QGroupBox)
from PySide6.QtCore import Qt
from utils.common.rutasarchivos import resource_path
from utils.common.metodosGenerales import MetodosGenerales
from controllers.ReporteController import ReporteController
from controllers.EmpresaController import EmpresaController
from utils.shared.actualizarReporte import ActualizarReporte
class ReporteAnexo1:
    doc = None
    
    # Función que justifica el párrafo en base a la longitud en caracteres
    def justificar_parrafo(doc, texto, max_caracteres_por_linea=80):
        # Separar el texto por saltos de línea
        lineas = texto.split('\n')
        
        for idx, linea in enumerate(lineas):
            parrafo = doc.add_paragraph(linea)
            
            # Calcular la longitud en caracteres de la línea
            longitud_linea = len(linea)
            
            # Justificar si la línea excede la longitud máxima en caracteres
            if longitud_linea > max_caracteres_por_linea:
                parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            else:
                parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    def agregar_parrafos_con_vinetas(doc, texto):
        # Dividir el texto en párrafos usando el salto de línea
        parrafos = texto.split('\n')
        
        for parrafo in parrafos:
            # Agregar cada párrafo con una viñeta
            nuevo_parrafo = doc.add_paragraph(parrafo, style='List Bullet')
            
            # Ajustar la sangría y alineación
            nuevo_parrafo.paragraph_format.left_indent = Inches(0.5)
            nuevo_parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def agregar_parrafos_alternando_negrita(cell, texto):
        # Dividir el texto en párrafos usando el salto de línea
        parrafos = texto.split('\n')

        # Eliminar cualquier párrafo vacío al principio
        parrafos = [p for p in parrafos if p.strip()]

        # Obtener el primer párrafo existente en la celda (si existe)
        if len(cell.paragraphs) > 0:
            nuevo_parrafo = cell.paragraphs[0]  # Usar el primer párrafo existente
        else:
            # Si no hay párrafos existentes, crear uno nuevo
            nuevo_parrafo = cell.add_paragraph()

        # Recorrer los párrafos divididos y agregar los textos con alternancia de estilos
        for i, parrafo in enumerate(parrafos):
            # Si no es el primer párrafo, agregar un salto de línea
            if i > 0:
                nuevo_parrafo.add_run("\n")

            # Eliminar los espacios innecesarios antes de agregar el texto
            run = nuevo_parrafo.add_run(parrafo.strip())
            
            # Alternar entre negrita y normal
            if i % 2 == 0:  # Para párrafos con índice par (negrita)
                run.bold = True
            else:  # Para párrafos con índice impar (normal)
                run.bold = False

    def set_cell_background_color(cell, color_hex):
        """Aplica un color de fondo en hexadecimal a una celda de la tabla."""
        cell_properties = cell._element.get_or_add_tcPr()
        cell_shading = OxmlElement('w:shd')
        cell_shading.set(qn('w:fill'), color_hex)  # Configura el color de fondo usando qn para el espacio de nombres
        cell_properties.append(cell_shading)
    
    def agregar_viñeta_numerica_negrita(doc, texto, estilo="Heading 2"):
        # Crear el párrafo con el estilo indicado
        parrafo = doc.add_paragraph(texto, style=estilo)

        # Aplicar formato al texto (negrita y tamaño de fuente)
        for run in parrafo.runs:
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0) 

        # Alineación (opcional)
        parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Configurar el párrafo para mantenerse junto con la siguiente tabla
        ReporteAnexo1.set_keep_together(parrafo)

        return parrafo
    
    def agregar_viñeta_letras_negrita(doc, texto, letra):
        # Crear un párrafo
        parrafo = doc.add_paragraph()

        # Agregar la viñeta en letra con formato en negrita
        run_letra = parrafo.add_run(f"{letra}. ")
        run_letra.font.bold = True  # Negrita
        run_letra.font.size = Pt(10)  # Tamaño de fuente ajustable

        # Agregar el texto principal en negrita
        run_texto = parrafo.add_run(texto)
        run_texto.font.bold = True  # Negrita
        run_texto.font.size = Pt(10)  # Tamaño de fuente ajustable

        # Alineación del párrafo
        parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Alineado a la izquierda
        
        # Configurar el párrafo para mantenerse junto con la siguiente tabla
        ReporteAnexo1.set_keep_together(parrafo)
    
    def reparar_imagen(img_stream):
        try:
            # Abrir la imagen usando PIL
            img = Image.open(img_stream)

            # Verificar las dimensiones de la imagen
            width, height = img.size
            if width == 0 or height == 0:
                raise ValueError("Dimensiones de la imagen no válidas")

            # Crear un nuevo BytesIO para guardar la imagen reparada
            repaired_img_stream = BytesIO()
            img.save(repaired_img_stream, format=img.format)
            repaired_img_stream.seek(0)

            return repaired_img_stream
        except Exception as e:
            raise ValueError(f"Error al reparar la imagen: {e}")

    # para que siempre este juntos algunos elementos
    def set_keep_together(parrafo):
        """
        Configura el párrafo para mantenerse junto con el siguiente elemento (tabla, párrafo, etc.).
        """
        p = parrafo._p  # Accede al elemento XML subyacente del párrafo
        pPr = p.get_or_add_pPr()
        keepNext = OxmlElement('w:keepNext')  # Mantener con el siguiente párrafo o elemento
        pPr.append(keepNext)
    
    def insertar_imagen_con_descripcion(doc, imagen_blob, descripcion, numero_figura):
        # Añadir la imagen centrada
        par_imagen = doc.add_paragraph()
        par_imagen.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_imagen = par_imagen.add_run()
        run_imagen.add_picture(imagen_blob, width=Inches(4))  # Ajustar tamaño según necesidad

        # Añadir la descripción centrada en negrita
        texto_figura = f"Figura N° {numero_figura}: {descripcion}"
        par_descripcion = doc.add_paragraph(texto_figura)
        par_descripcion.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_descripcion = par_descripcion.runs[0]
        run_descripcion.bold = True
        run_descripcion.font.size = Pt(10)  # Tamaño de letra ajustable
        
    # Crear un documento nuevo
    def crear_documento():
        return Document()

    # Configurar página en tamaño A4 y aplicar márgenes APA
    def configurar_pagina(section):
        section.page_height = Inches(11.69)  # Tamaño A4 en pulgadas (29.7 cm)
        section.page_width = Inches(8.27)    # Tamaño A4 en pulgadas (21 cm)
        section.top_margin = Inches(1.0)     # Margen superior APA (1 pulgada)
        section.bottom_margin = Inches(1.0)  # Margen inferior APA (1 pulgada)
        section.left_margin = Inches(1.0)    # Margen izquierdo APA (1 pulgada)
        section.right_margin = Inches(1.0)   # Margen derecho APA (1 pulgada)

    # Configurar estilo de fuente y tamaño según APA
    def configurar_fuente(doc):
        style = doc.styles['Normal']
        style.font.name = 'Arial'  # Fuente APA estándar
        style.font.size = Pt(12) # Tamaño de fuente APA
        
    # Función para crear un título con formato Heading 1 y espaciado automático
    def agregar_titulo_1(doc, texto, alineacion=WD_PARAGRAPH_ALIGNMENT.CENTER):
        # Agregar un párrafo vacío antes del título para espacio adicional

        # Crear el párrafo del título y aplicar el estilo y formato
        titulo = doc.add_paragraph(texto)
        titulo.style = 'Heading 1'
        titulo.alignment = alineacion  # Aplicar la alineación directamente

        # Aplicar formato específico: Arial, tamaño, color, negrita
        run = titulo.runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(12)  # Tamaño APA, ajusta según la norma
        run.font.color.rgb = RGBColor(0, 0, 0)  # Color negro
        run.bold = True  # Negrita si es requerido

        # Ajustar espacios antes y después del título
        titulo.paragraph_format.space_before = Pt(0)  # Sin espacio extra antes
        titulo.paragraph_format.space_after = Pt(0)  # Espacio después del título

        return titulo

    # Función para crear un título con formato Heading 2
    def agregar_titulo_2(doc, texto, alineacion=WD_PARAGRAPH_ALIGNMENT.LEFT):
        # Crear el párrafo del título y aplicar el estilo y formato
        titulo = doc.add_paragraph(texto)
        titulo.style = 'Heading 2'
        titulo.alignment = alineacion  # Aplicar la alineación directamente

        # Aplicar formato específico: Arial, tamaño, color, sin ajuste de espacios
        run = titulo.runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(12)  # Tamaño para título 2, ajusta si es necesario
        run.font.color.rgb = RGBColor(0, 0, 0)  # Color negro
        run.bold = True  # Negrita si es requerido

        return titulo
    
    # Añadir encabezado en la segunda sección del documento
    def agregar_encabezado(doc, datos_generales):
        if datos_generales:
            logotipo_blob = datos_generales[5]
            if len(doc.sections) > 1:
                section = doc.sections[1]
            else:
                section = doc.sections[0]
            header = section.header
            header.is_linked_to_previous = False
            header_table = header.add_table(rows=1, cols=2, width=Inches(7))
            header_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for cell in header_table.rows[0].cells:
                cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            cell_logo = header_table.cell(0, 0)
            if logotipo_blob:
                try:
                    # Convertir el blob a un archivo en memoria y agregarlo como imagen
                    image_stream = BytesIO(logotipo_blob)
                    run = cell_logo.paragraphs[0].add_run()
                    # Ajustar el tamaño de la imagen al ancho deseado
                    run.add_picture(image_stream, width=Inches(1.5))
                except ZeroDivisionError:
                    try:
                        img_stream = ReporteAnexo1.reparar_imagen(image_stream)
                        # Ajustar el tamaño de la imagen al ancho deseado
                        run.add_picture(img_stream, width=Inches(1.5))
                    except Exception as e:
                        cell_logo.text = f"Error al cargar el logotipo: {str(e)}"
                except Exception as e:
                    cell_logo.text = f"Error al cargar el logotipo: {str(e)}"
                        
            else:
                cell_logo.text = " "

            cell_text = header_table.cell(0, 1)
            header_text = cell_text.paragraphs[0]
            header_text.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            run_text = header_text.add_run(f"{datos_generales[1]}\nJefatura de Geotecnia")
            run_text.font.size = Pt(10)
            run_text.bold = True

            tbl_pr = header_table._element.find(qn("w:tblPr"))
            if tbl_pr is None:
                tbl_pr = OxmlElement("w:tblPr")
                header_table._element.insert(0, tbl_pr)
            tbl_borders = OxmlElement("w:tblBorders")
            bottom_border = OxmlElement("w:bottom")
            bottom_border.set(qn("w:val"), "single")
            bottom_border.set(qn("w:sz"), "15")
            bottom_border.set(qn("w:space"), "0")
            bottom_border.set(qn("w:color"), "263656")
            tbl_borders.append(bottom_border)
            tbl_pr.append(tbl_borders)

            # Agregar un párrafo vacío debajo de la tabla
            header.add_paragraph("")

    # Configurar pie de página con borde superior y número de página en la segunda sección
    def agregar_pie_pagina(doc):
        if len(doc.sections) > 1:
            section = doc.sections[1]
        else:
            section = doc.sections[0]
        section.footer.is_linked_to_previous = False
        footer = section.footer
        footer_table = footer.add_table(rows=1, cols=1, width=Inches(7))
        footer_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_cell = footer_table.cell(0, 0)

        page_number_paragraph = footer_cell.paragraphs[0]
        page_number_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        page_num_run = page_number_paragraph.add_run()
        page_num_run.font.size = Pt(10)

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        page_num_run._r.append(fldChar)

        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE \\* Arabic \\* MERGEFORMAT"
        page_num_run._r.append(instrText)

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "end")
        page_num_run._r.append(fldChar)

        tbl_pr_footer = footer_table._element.find(qn("w:tblPr"))
        if tbl_pr_footer is None:
            tbl_pr_footer = OxmlElement("w:tblPr")
            footer_table._element.insert(0, tbl_pr_footer)
        tbl_borders_footer = OxmlElement("w:tblBorders")
        top_border = OxmlElement("w:top")
        top_border.set(qn("w:val"), "single")
        top_border.set(qn("w:sz"), "15")
        top_border.set(qn("w:space"), "0")
        top_border.set(qn("w:color"), "263656")
        tbl_borders_footer.append(top_border)
        tbl_pr_footer.append(tbl_borders_footer)

        # Función para agregar caratula sin pie ni encabezado
    
    # Función para agregar caratula sin pie ni encabezado
    def agregar_caratula(doc, logotipo_blob, datos, mes_anio):
        # Acceder a la primera sección del documento
        if datos:
            section = doc.sections[0]
            ReporteAnexo1.configurar_pagina(section)
            section.different_first_page_header_footer = True  # Sin encabezado ni pie en la primera página

            # Añadir espacio vacío
            for _ in range(5):
                doc.add_paragraph("")

            # Si se pasa un BLOB de la imagen, usar BytesIO para insertarlo
            if logotipo_blob:
                try:
                    # Crear un objeto BytesIO a partir del BLOB
                    image_stream  = BytesIO(logotipo_blob)

                    # Intentar insertar la imagen en la carátula
                    img_paragraph = doc.add_paragraph()
                    img_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = img_paragraph.add_run()
                    run.add_picture(image_stream , width=Inches(2.5))
                except ZeroDivisionError:
                    # Si hay un error de división por cero, intentar reparar la imagen
                    try:
                        img_stream = ReporteAnexo1.reparar_imagen(image_stream)
                        run.add_picture(img_stream, width=Inches(2.5))
                    except Exception as e:
                        doc.add_paragraph(f"Error al insertar la imagen: {e}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except Exception as e:
                    doc.add_paragraph(f"Error al insertar la imagen: {e}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                doc.add_paragraph("No se proporcionó imagen.").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Título del reporte
            title = doc.add_paragraph(datos[2])
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title.runs[0].bold = True
            title.runs[0].font.size = Pt(20)
            title.add_run("\n")

            # Subtítulo del reporte
            subtitle = doc.add_paragraph(f"{datos[3]}\nMES: {mes_anio}")
            subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            subtitle.runs[0].font.size = Pt(18)

            # Añadir más espacio vacío
            for _ in range(6):
                doc.add_paragraph("")

            # Pie de página
            footer_paragraph = doc.add_paragraph(f"{datos[4]}\nElaborado por: {datos[5]}")
            footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            footer_paragraph.runs[0].bold = True
            footer_paragraph.runs[0].font.size = Pt(11)

            # Crear una nueva sección para la siguiente página
            new_section = doc.add_section(WD_SECTION.NEW_PAGE)
            ReporteAnexo1.configurar_pagina(new_section)
            new_section.different_first_page_header_footer = False  # Permitir encabezado y pie
    
    # Añadir el contenido del "MEMORANDUM" en la segunda página utilizando tabla para alinear texto
    def agregar_contenido_memorandum(doc, datos, responsable, unidad_minera, dia_mes_anio):
        if datos:
            invisible_space = "\u00A0" * 10
            firma_blob, titulomemo, codigomemo, destinatario, remitente, asunto, descrireporte = None, "MEMORANDUN", "", "", "", "", ""
            if responsable:
                firma_blob = responsable[6]
            if datos:
                titulomemo, codigomemo, destinatario, remitente, asunto, descrireporte = datos[6], datos[7], datos[8], datos[9], datos[10], datos[11]
            # Título del Docuemnto
            titulo = doc.add_paragraph(titulomemo)
            titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            titulo_run = titulo.runs[0]
            titulo_run.font.size = Pt(14)
            titulo_run.bold = True
            # CODIGO
            referencia = doc.add_paragraph(codigomemo)
            referencia.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            referencia.runs[0].font.size = Pt(10)
            referencia.runs[0].font.bold = True 

            # Tabla invisible para alinear texto de forma organizada
            table = doc.add_table(rows=4, cols=2)
            table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            table.autofit = False
            table.columns[0].width = Inches(1.5)
            table.columns[1].width = Inches(5.0)

            # Reducir los márgenes internos de las celdas para minimizar el espacio entre columnas
            for row in table.rows:
                for cell in row.cells:
                    tc_pr = cell._element.get_or_add_tcPr()
                    tc_mar = OxmlElement('w:tcMar')
                    for side in ['top', 'left', 'bottom', 'right']:
                        margin = OxmlElement(f'w:{side}')
                        margin.set(qn('w:w'), '0')  # Sin margen
                        margin.set(qn('w:type'), 'dxa')
                        tc_mar.append(margin)
                    tc_pr.append(tc_mar)

            # Rellenar la tabla con texto ajustado
            cell_0_0 = table.cell(0, 0).paragraphs[0].add_run(f"Para:{invisible_space}")
            cell_0_0.bold = True
            # Añadir texto alternando negrita y normal a la segunda celda de la fila
            cell_0_1 = table.cell(0, 1)
            ReporteAnexo1.agregar_parrafos_alternando_negrita(cell_0_1, destinatario)

            cell_1_0 = table.cell(1, 0).paragraphs[0].add_run("De: ")
            cell_1_0.bold = True
            cell_1_1 = table.cell(1, 1)
            ReporteAnexo1.agregar_parrafos_alternando_negrita(cell_1_1, remitente)

            cell_2_0 = table.cell(2, 0).paragraphs[0].add_run("Fecha: ")
            cell_2_0.bold = True
            cell_2_1 = table.cell(2, 1).paragraphs[0]
            cell_2_1.add_run(dia_mes_anio)

            cell_3_0 = table.cell(3, 0).paragraphs[0].add_run("Asunto: ")
            cell_3_0.bold = True
            cell_3_1 = table.cell(3, 1).paragraphs[0]
            cell_3_1.add_run(asunto)

            # Añadir el cuerpo del memorandum
            ReporteAnexo1.justificar_parrafo(doc, descrireporte)

            # Añadir espacio antes de la firma
            for _ in range(2):
                doc.add_paragraph("")

            # Añadir la firma y detalles
            if firma_blob:                
                try:
                    # Convertir el blob a un archivo en memoria y agregarlo como imagen
                    firma_stream = BytesIO(firma_blob)
                    # Insertar la firma en el documento
                    firma_paragraph = doc.add_paragraph()
                    firma_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    firma_run = firma_paragraph.add_run()
                    firma_run.add_picture(firma_stream, width=Inches(1.5))
                except ZeroDivisionError:
                    try:
                        fir_stream = ReporteAnexo1.reparar_imagen(firma_stream)
                        # Insertar la firma en el documento
                        firma_paragraph = doc.add_paragraph()
                        firma_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        firma_run = firma_paragraph.add_run()
                        firma_run.add_picture(fir_stream, width=Inches(1.5))
                    except Exception as e:
                        doc.add_paragraph(f"Error al insertar la imagen: {e}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                except Exception as e:
                    doc.add_paragraph(f"Error al insertar la firma: {e}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                doc.add_paragraph("Firma no proporcionada.").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Información adicional sobre la firma
            if responsable:
                firma_paragraph = doc.add_paragraph(f"{responsable[2]}\n{responsable[3]}\n{unidad_minera}")
                firma_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Insertar salto de página
            doc.add_page_break()

    # Función para agregar la tercera página con la imagen y los títulos
    def agregar_tercera_pagina(doc, respuesta, mes_anio):
        if respuesta:
            # tipo reporte
            titulo = doc.add_paragraph(respuesta[12])
            titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            titulo.runs[0].bold = True
            titulo.runs[0].font.size = Pt(14)
            # fecha
            subtitulo = doc.add_paragraph(mes_anio)
            subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            subtitulo.runs[0].bold = True
            subtitulo.runs[0].font.size = Pt(12)
            # componente
            componente = doc.add_paragraph(respuesta[13])
            componente.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            componente.runs[0].bold = True
            componente.runs[0].font.size = Pt(12)
            doc.add_paragraph("")
            # Insertar la imagen BLOB en el centro de la página
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            MAX_WIDTH = 6.0  # pulgadas
            MAX_HEIGHT = 7.0  # pulgadas
            if respuesta[14]:
                try:
                    image_stream = BytesIO(respuesta[14])
                    # Abrir imagen con PIL para obtener dimensiones
                    image = Image.open(image_stream)
                    width_px, height_px = image.size
                    dpi = image.info.get('dpi', (96, 96))[0]  # usar 96 dpi por defecto
                    width_in = width_px / dpi
                    height_in = height_px / dpi
                    # Escalar proporcionalmente para ajustarse a los límites máximos
                    scale = min(MAX_WIDTH / width_in, MAX_HEIGHT / height_in)
                    final_width = Inches(width_in * scale)
                    final_height = Inches(height_in * scale)
                    # Reiniciar el stream para Word (PIL lo ha movido)
                    image_stream.seek(0)
                    run.add_picture(image_stream, width=final_width, height=final_height)
                except ZeroDivisionError:
                    try:
                        img_stream = ReporteAnexo1.reparar_imagen(image_stream)
                        run.add_picture(img_stream, width=Inches(6.5))
                    except Exception as e:
                        paragraph.add_run(f"Error al insertar la imagen: {str(e)}")
                except Exception as e:
                    paragraph.add_run(f"Error al insertar la imagen: {str(e)}")
            # Insertar salto de página
            doc.add_page_break()
    
    # Función para agregar la página del índice con marcador
    def agregar_indice_con_marcador(doc):
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        ReporteAnexo1.configurar_pagina(new_section)
        # Añadir el título "Índice" centrado y formateado
        titulo = doc.add_paragraph("ÍNDICE")
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = titulo.runs[0]
        run.bold = True
        run.font.size = Pt(14)
        run.font.underline = False  # Sin subrayado
        run.font.color.rgb = RGBColor(0, 0, 0)  # Color negro
        
        # Insertar marcador donde se añadirá la tabla de contenido
        toc_placeholder = doc.add_paragraph()
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), "0")
        bookmark_start.set(qn("w:name"), "TOCPlaceholder")
        toc_placeholder._p.append(bookmark_start)
        
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), "0")
        toc_placeholder._p.append(bookmark_end)
        doc.add_page_break()

    # Función para agregar el contenido de la quinta página con ajustes adicionales y alineación justificada
    def agregar_contenido_quinta_pagina(doc,respuesta):  
        if respuesta:  
            ReporteAnexo1.agregar_titulo_1(doc,"INTRODUCCIÓN",WD_PARAGRAPH_ALIGNMENT.LEFT)
            # Objetivo
            ReporteAnexo1.agregar_titulo_1(doc,'OBJETIVO',WD_PARAGRAPH_ALIGNMENT.LEFT)
            ReporteAnexo1.justificar_parrafo(doc,respuesta[15])
            # Finalidad
            ReporteAnexo1.agregar_titulo_1(doc,'FINALIDAD',WD_PARAGRAPH_ALIGNMENT.LEFT)
            ReporteAnexo1.justificar_parrafo(doc,respuesta[16])

            # Ámbito de Aplicación
            ReporteAnexo1.agregar_titulo_1(doc,'ÁMBITO DE APLICACIÓN',WD_PARAGRAPH_ALIGNMENT.LEFT)
            ReporteAnexo1.justificar_parrafo(doc,respuesta[17])

            # Reporte Mensual con viñetas y justificado
            ReporteAnexo1.agregar_titulo_1(doc,'REPORTE MENSUAL',WD_PARAGRAPH_ALIGNMENT.LEFT)
            ReporteAnexo1.agregar_parrafos_con_vinetas(doc,respuesta[18])

            if respuesta[19] and respuesta[19].strip():
                anexo1 = doc.add_paragraph(f"ANEXO 1: {respuesta[19]}")
                anexo1.style = 'Normal'
                anexo1.paragraph_format.left_indent = Inches(0.5)
                anexo1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            # Insertar salto de página
            doc.add_page_break()
    
    ################## ANEXO 1 #####################
    def generar_tabla_resumen_ejecutivo(doc, componentes):
        if componentes:
            idcomponentes = [compo[0] for compo in componentes]
            respuesta = ReporteController.ctrlObtenerTablaResumenEjecutivoAnexo1(idcomponentes)
            if respuesta:
                # Inicializar el contador de letras (a., b., c., ...)
                letra_actual = 97  # Código ASCII de 'a'
                for fila in respuesta:
                    # Si hay más de una fila, añadir un título con viñeta alfabética
                    if len(respuesta) > 1:
                        letra = chr(letra_actual)
                        ReporteAnexo1.agregar_viñeta_letras_negrita(doc, fila[0], letra)
                        letra_actual += 1  # Incrementar letra para la siguiente tabla

                    # Crear una nueva tabla con 5 filas y 12 columnas
                    table = doc.add_table(rows=5, cols=12)
                    table.style = 'Table Grid'
                    table.allow_autofit = True

                    # Fila 1: DEPÓSITO - 6 columnas combinadas cada celda
                    cell_00 = table.cell(0, 0)
                    cell_00.text = fila[4]  # DEPÓSITO:
                    ReporteAnexo1.set_cell_background_color(table.cell(0, 0), "F0F0F0")
                    cell_00.merge(table.cell(0, 5))  # Combinar primeras 6 columnas
                    cell_01 = table.cell(0, 6)
                    cell_01.text = fila[5]  # DEPÓSITO DE RELAVE
                    cell_01.merge(table.cell(0, 11))  # Combinar últimas 6 columnas

                    # Fila 2: ÚLTIMA AUTORIZACIÓN DE FUNCIONAMIENTO
                    cell_10 = table.cell(1, 0)
                    cell_10.text = fila[6]  # ÚLTIMA AUTORIZACIÓN DE FUNCIONAMIENTO:
                    ReporteAnexo1.set_cell_background_color(table.cell(1, 0), "F0F0F0")
                    cell_10.merge(table.cell(1, 5))  # Combinar primeras 6 columnas
                    cell_11 = table.cell(1, 6)
                    cell_11.text = fila[7]  # PAMA
                    cell_11.merge(table.cell(1, 11))  # Combinar últimas 6 columnas

                    # Fila 3: MES / AÑO
                    cell_20 = table.cell(2, 0)
                    cell_20.text = fila[8]  # MES / AÑO:
                    ReporteAnexo1.set_cell_background_color(table.cell(2, 0), "F0F0F0")
                    cell_20.merge(table.cell(2, 5))  # Combinar primeras 6 columnas
                    cell_21 = table.cell(2, 6)
                    cell_21.text = fila[9]  # SETIEMBRE 2024
                    cell_21.merge(table.cell(2, 11))  # Combinar últimas 6 columnas

                    # Fila 4: Declaración (primera) - 8 columnas combinadas + 4 individuales
                    cell_30 = table.cell(3, 0)
                    cell_30.text = fila[10]  # Primera declaración
                    cell_30.merge(table.cell(3, 7))  # Combinar las primeras 8 columnas

                    table.cell(3, 8).text = "SI"
                    table.cell(3, 9).text = "X" if fila[11] == "SI" else ""
                    table.cell(3, 10).text = "NO"
                    table.cell(3, 11).text = "X" if fila[11] == "NO" else ""

                    # Aplicar color a las celdas con "X" en SI o NO
                    if fila[11] == "SI":
                        ReporteAnexo1.set_cell_background_color(table.cell(3, 8), "C6E0B4")
                        ReporteAnexo1.set_cell_background_color(table.cell(3, 9), "C6E0B4")
                    elif fila[11] == "NO":
                        ReporteAnexo1.set_cell_background_color(table.cell(3, 10), "C6E0B4")
                        ReporteAnexo1.set_cell_background_color(table.cell(3, 11), "C6E0B4")

                    # Fila 5: Declaración (segunda) - 8 columnas combinadas + 4 individuales
                    cell_40 = table.cell(4, 0)
                    cell_40.text = fila[12]  # Segunda declaración
                    cell_40.merge(table.cell(4, 7))  # Combinar las primeras 8 columnas

                    table.cell(4, 8).text = "SI"
                    table.cell(4, 9).text = "X" if fila[13] == "SI" else ""
                    table.cell(4, 10).text = "NO"
                    table.cell(4, 11).text = "X" if fila[13] == "NO" else ""

                    # Aplicar color a las celdas con "X" en SI o NO
                    if fila[13] == "SI":
                        ReporteAnexo1.set_cell_background_color(table.cell(4, 8), "C6E0B4")
                        ReporteAnexo1.set_cell_background_color(table.cell(4, 9), "C6E0B4")
                    elif fila[13] == "NO":
                        ReporteAnexo1.set_cell_background_color(table.cell(4, 10), "C6E0B4")
                        ReporteAnexo1.set_cell_background_color(table.cell(4, 11), "C6E0B4")

                    # Aplicar tamaño de texto global (10 puntos) a todas las celdas
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.runs[0].font.size = Pt(10)
                                
                    # Aplicar estilos a las filas 4 y 5
                    for row_index in [3, 4]:  # Filas 4 y 5
                        for cell in table.row_cells(row_index)[8:12]:  # Columnas SI, X, NO, X
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    for cell in [cell_30, cell_40]:  # Declaraciones de texto combinado
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            paragraph.runs[0].font.size = Pt(10)

                    # Centramos la tabla y agregamos espacio al final
                    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Separar las tablas con un espacio
                    doc.add_paragraph("")
    
    def generar_tabla_parametros(doc, componentes):
        if componentes:
            idcomponentes = [compo[0] for compo in componentes]
            respuesta = ReporteController.ctrlObtenerParametrosAnexo1(idcomponentes)
            if respuesta:
                letra = "a"
                grupos = {}
                for fila in respuesta:
                    grupo = fila[2]  # Usamos el idcomponente
                    if grupo not in grupos:
                        grupos[grupo] = []
                    grupos[grupo].append(fila)
                # Verificar si hay más de un grupo
                agregar_letra = len(grupos) > 1

                # Iterar sobre los grupos y generar las tablas
                for grupo, filas in grupos.items():
                    # Agregar el título con la letra correspondiente en minúscula si hay más de un grupo
                    if agregar_letra:
                        titulo = filas[0][0]
                        ReporteAnexo1.agregar_viñeta_letras_negrita(doc, titulo, letra)
                        letra = chr(ord(letra) + 1)  # Incrementar la letra para la siguiente tabla

                    # Crear la tabla con 1 fila para la cabecera y 8 columnas
                    table = doc.add_table(rows=1, cols=8)  # 8 columnas
                    table.style = 'Table Grid'

                    # Configuración de las cabeceras
                    table.cell(0, 0).text = "Descripción"
                    table.cell(0, 1).text = "Parámetros"
                    table.cell(0, 3).text = "Condición Actual"
                    table.cell(0, 7).text = "Comentarios"
                    # Pintar el encabezado completo
                    for cell in table.rows[0].cells:
                        ReporteAnexo1.set_cell_background_color(cell, "F0F0F0")

                    # Combinar celdas para las cabeceras
                    table.cell(0, 1).merge(table.cell(0, 2))  # "Parámetros" abarca 2 columnas
                    table.cell(0, 3).merge(table.cell(0, 6))  # "Condición Actual" abarca 4 columnas

                    # Alineación de las celdas de cabecera
                    for cell in table.rows[0].cells:
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            paragraph.runs[0].font.size = Pt(10)
                            paragraph.runs[0].font.bold = True

                    # Insertar datos en las filas
                    for row_data in filas:
                        row = table.add_row().cells  # Agregar una nueva fila y obtener sus celdas
                        row[0].text = row_data[3]  # Descripción
                        row[1].text = row_data[4]  # Parámetro 1
                        row[2].text = row_data[5]  # Parámetro 2

                        # "Condición Actual" - Cumple y No Cumple
                        row[3].text = "Cumple"
                        row[4].text = "(X)" if row_data[6] == "CUMPLE" else "(\u00A0)"
                        row[5].text = "No Cumple"
                        row[6].text = "(X)" if row_data[6] == "NO CUMPLE" else "(\u00A0)"

                        # Pintar celdas según la condición
                        if row_data[6] == "CUMPLE":
                            ReporteAnexo1.set_cell_background_color(row[3], "C6E0B4")  # Pintar columna "Cumple"
                            ReporteAnexo1.set_cell_background_color(row[4], "C6E0B4")  # Pintar celda "X" en "Cumple"
                        elif row_data[6] == "NO CUMPLE":
                            ReporteAnexo1.set_cell_background_color(row[5], "C6E0B4")  # Pintar columna "No Cumple"
                            ReporteAnexo1.set_cell_background_color(row[6], "C6E0B4")  # Pintar celda "X" en "No Cumple"

                        # Comentarios
                        row[7].text = row_data[7]  # Comentarios

                        # Alinear texto en las celdas de contenido
                        for j in range(8):
                            cell = row[j]
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = (
                                    WD_PARAGRAPH_ALIGNMENT.CENTER if j in [3, 4, 5, 6] else WD_PARAGRAPH_ALIGNMENT.LEFT
                                )
                                paragraph.runs[0].font.size = Pt(10)

                    # Centramos la tabla y agregamos espacio al final
                    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Separar las tablas con un espacio
                    doc.add_paragraph("")
    
    def generar_tabla_condiciones_fisicas(doc, componentes):
        if componentes:
            idcomponentes = [compo[0] for compo in componentes]
            respuesta = ReporteController.ctrlObtenerCondicionesFisicasAnexo1(idcomponentes)
            if respuesta:
                letra = "a"
                grupos = {}
                for fila in respuesta:
                    grupo = fila[2]  # Usamos el idcomponente
                    if grupo not in grupos:
                        grupos[grupo] = []
                    grupos[grupo].append(fila)

                # Verificar si hay más de un grupo
                if len(grupos) > 1:
                    agregar_letra = True
                else:
                    agregar_letra = False

                # Iterar sobre los grupos y generar las tablas
                for grupo, filas in grupos.items():
                    # Solo agregar la letra si hay más de un grupo
                    if agregar_letra:
                        titulo=filas[0][0]
                        # Agregar el título con la letra correspondiente en minúscula
                        ReporteAnexo1.agregar_viñeta_letras_negrita(doc, titulo, letra)
                        letra = chr(ord(letra) + 1)  # a -> b -> c -> ...

                    # Crear la tabla con 1 fila para la cabecera y 6 columnas
                    table = doc.add_table(rows=1, cols=6)
                    table.style = 'Table Grid'  # Estilo de tabla

                    # Configuración de las celdas de encabezado
                    table.cell(0, 0).text = "Condiciones de los taludes"
                    table.cell(0, 1).text = "Condición Actual"
                    table.cell(0, 5).text = "Comentarios"
                    # Pintar el encabezado completo
                    for cell in table.rows[0].cells:
                        ReporteAnexo1.set_cell_background_color(cell, "F0F0F0")

                    # Combinar celdas para que "Condición Actual" abarque cuatro columnas
                    table.cell(0, 1).merge(table.cell(0, 4))

                    # Alineación y tamaño de fuente en las celdas de encabezado
                    for cell in [table.cell(0, 0), table.cell(0, 1), table.cell(0, 5)]:
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            if paragraph.runs:
                                paragraph.runs[0].font.size = Pt(10)

                    # Iterar sobre los datos y agregar filas
                    for row_data in filas:
                        row = table.add_row().cells  # Crear una nueva fila y obtener sus celdas
                        row[0].text = row_data[3]  # Condición (Columna 0)
                        row[5].text = row_data[5]  # Comentarios (Columna 5)

                        # Verificar el tipo de condición (columna 7)
                        tipo_condicion = row_data[6]  # Tipo de condición

                        if tipo_condicion == 3:  # Tipo 3 es solo texto, combinamos celdas
                            # Combinar las celdas de "Condición Actual" (Columnas 1, 2, 3, 4)
                            row[1].merge(row[2]).merge(row[3]).merge(row[4])

                            # Si la condición está vacía, asignar "N.A"
                            condicion = row_data[4] if row_data[4] else "N. A"
                            row[1].text = condicion  # Agregar el texto (comentario o estado)
                            
                            # Aplicar fondo o estilo si es necesario
                            ReporteAnexo1.set_cell_background_color(row[1], "C6E0B4")  # Ejemplo de fondo de color

                        elif tipo_condicion == 1:  # Tipo 1: "Cumple" o "No Cumple"
                            row[1].text = "CUMPLE"  # Columna 1: "Cumple"
                            row[2].text = "(X)" if row_data[4] == "CUMPLE" else "()"  # Columna 2: "X" para cumple
                            row[3].text = "NO CUMPLE"  # Columna 3: "No Cumple"
                            row[4].text = "(X)" if row_data[4] == "NO" else "()"  # Columna 4: "X" para no cumple
                            row[5].text = row_data[5]  # Comentarios

                            # Aplicar color de fondo solo si la X está presente
                            if row_data[4] == "CUMPLE":
                                ReporteAnexo1.set_cell_background_color(row[1], "C6E0B4")  # Color verde para "Cumple"
                                ReporteAnexo1.set_cell_background_color(row[2], "C6E0B4")  # Pintar la X de "Cumple"
                            if row_data[4] == "NO":
                                ReporteAnexo1.set_cell_background_color(row[3], "C6E0B4")  # Color rojo para "No Cumple"
                                ReporteAnexo1.set_cell_background_color(row[4], "C6E0B4")  # Pintar la X de "No Cumple"

                        elif tipo_condicion == 2:  # Tipo 2: "SI" o "NO"
                            row[1].text = "SI"  # Columna 1: "SI"
                            row[2].text = "(X)" if row_data[4] == "SI" else "()"  # Columna 2: "X" para SI
                            row[3].text = "NO"  # Columna 3: "NO"
                            row[4].text = "(X)" if row_data[4] == "NO" else "()"  # Columna 4: "X" para NO
                            row[5].text = row_data[5]  # Comentarios

                            # Aplicar color de fondo solo si la X está presente
                            if row_data[4] == "SI":
                                ReporteAnexo1.set_cell_background_color(row[1], "C6E0B4")  # Color verde para "SI"
                                ReporteAnexo1.set_cell_background_color(row[2], "C6E0B4")  # Pintar la X de "SI"
                            if row_data[4] == "NO":
                                ReporteAnexo1.set_cell_background_color(row[3], "C6E0B4")  # Color rojo para "NO"
                                ReporteAnexo1.set_cell_background_color(row[4], "C6E0B4")  # Pintar la X de "NO"

                        # Alinear y ajustar el tamaño de la fuente en cada celda
                        for j in range(6):
                            cell = row[j]
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            for paragraph in cell.paragraphs:
                                # Centramos las columnas de "Condición Actual" y "Comentarios"
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if j in [1, 2, 3, 4] else WD_PARAGRAPH_ALIGNMENT.LEFT
                                if paragraph.runs:
                                    paragraph.runs[0].font.size = Pt(10)

                    # Centramos la tabla y agregamos espacio al final
                    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Separar las tablas con un espacio
                    doc.add_paragraph("")
    
    def generar_tabla_operatividad_equipos(doc, componentes):
        if componentes:
            idcomponentes = [compo[0] for compo in componentes]
            respuesta = ReporteController.ctrlObtenerOperatividadEquipos(idcomponentes)
            if respuesta:
                letra = "a"
                grupos = {}
                for fila in respuesta:
                    grupo = fila[2]  # Usamos la columna idcomponente
                    if grupo not in grupos:
                        grupos[grupo] = []
                    grupos[grupo].append(fila)

                # Verificar si hay más de un grupo
                if len(grupos) > 1:
                    agregar_letra = True
                else:
                    agregar_letra = False

                # Iterar sobre los grupos y generar las tablas
                for grupo, filas in grupos.items():
                    # Solo agregar la letra si hay más de un grupo
                    if agregar_letra:
                        titulo=filas[0][0]
                        ReporteAnexo1.agregar_viñeta_letras_negrita(doc, titulo, letra)
                        letra = chr(ord(letra) + 1)  # a -> b -> c -> ...

                    # Crear la tabla con 2 filas para encabezados y 6 columnas
                    table = doc.add_table(rows=2, cols=6)
                    table.style = 'Table Grid'

                    # Configuración de las celdas de encabezado
                    table.cell(0, 0).text = "Instrumentación"
                    table.cell(0, 1).text = "Condición Actual"
                    table.cell(0, 3).text = "Performance"
                    table.cell(0, 5).text = "Comentario"
                    # Pintar el encabezado completo
                    for cell in table.rows[0].cells:
                        ReporteAnexo1.set_cell_background_color(cell, "F0F0F0")

                    # Combinar celdas para los encabezados
                    table.cell(0, 1).merge(table.cell(0, 2))  # Condición Actual
                    table.cell(0, 3).merge(table.cell(0, 4))  # Performance
                    table.cell(0, 5).merge(table.cell(1, 5))  # Comentario
                    table.cell(0, 0).merge(table.cell(1, 0))  # Instrumentación

                    # Subcategorías de encabezado
                    table.cell(1, 1).text = "SI"
                    table.cell(1, 2).text = "NO"
                    table.cell(1, 3).text = "Cantidad"
                    table.cell(1, 4).text = "Operativo"

                    # Estilo de las celdas de encabezado
                    for row in table.rows:
                        for cell in row.cells:
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                if paragraph.runs:
                                    paragraph.runs[0].font.size = Pt(10)

                    # Procesar los datos y llenar las filas de la tabla
                    for data in filas:
                        # Extraer los valores necesarios
                        descripcion = data[3]  # Descripción o Instrumentación
                        condicion = data[4]    # Condición Actual (SI/NO)
                        comentario = data[7]   # Comentario
                        cantidad = data[5]       # Puedes ajustar este valor según corresponda
                        operativo = "SI" if data[6] == "SI" else "NO"

                        # Crear una nueva fila
                        row = table.add_row().cells
                        row[0].text = descripcion  # Instrumentación
                        row[1].text = "(X)" if condicion == "SI" else "()"  # SI
                        row[2].text = "(X)" if condicion == "NO" else "()"  # NO
                        row[3].text = cantidad  # Cantidad
                        row[4].text = operativo  # Operativo
                        row[5].text = comentario  # Comentario

                        # Aplicar color de fondo a las celdas de SI/NO
                        if "(X)" in row[1].text:  # SI
                            ReporteAnexo1.set_cell_background_color(row[1], "C6E0B4")
                        if "(X)" in row[2].text:  # NO
                            ReporteAnexo1.set_cell_background_color(row[2], "C6E0B4")

                        # Ajustar alineación y tamaño de fuente
                        for j in range(6):
                            cell = row[j]
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            for paragraph in cell.paragraphs:
                                if j == 5:  # Comentario
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                else:
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                if paragraph.runs:
                                    paragraph.runs[0].font.size = Pt(10)

                    # Centramos la tabla y agregamos espacio al final
                    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    doc.add_paragraph("")
    
    def generar_tabla_observaciones(doc, componentes):
        if componentes:
            idcomponentes = [compo[0] for compo in componentes]
            respuesta = ReporteController.ctrlObtenerObservacionesAnexo1(idcomponentes)
            if respuesta:
                letra = "a"
                grupos = {}
                for fila in respuesta:
                    grupo = fila[2]  # Usar la columna idcomponente
                    if grupo not in grupos:
                        grupos[grupo] = []
                    grupos[grupo].append(fila)

                # Verificar si hay más de un grupo
                agregar_letra = len(grupos) > 1

                # Iterar sobre los grupos y generar tablas para cada grupo
                for grupo, filas in grupos.items():
                    # Si hay varios grupos, agregar una viñeta con la letra correspondiente
                    if agregar_letra:
                        titulo=filas[0][0]
                        ReporteAnexo1.agregar_viñeta_letras_negrita(doc, titulo, letra)
                        letra = chr(ord(letra) + 1)  # Pasar a la siguiente letra

                    # Crear la tabla con 1 fila para la cabecera y 6 columnas
                    table = doc.add_table(rows=1, cols=6)
                    table.style = 'Table Grid'

                    # Configuración de las celdas de encabezado
                    table.cell(0, 0).text = "Descripción"
                    table.cell(0, 1).text = "Condición Actual"
                    table.cell(0, 5).text = "Comentarios"
                    # Pintar el encabezado completo
                    for cell in table.rows[0].cells:
                        ReporteAnexo1.set_cell_background_color(cell, "F0F0F0")
                    # Combinar las celdas de "Condición Actual"
                    table.cell(0, 1).merge(table.cell(0, 4))

                    # Alinear y ajustar las celdas de encabezado
                    for cell in table.rows[0].cells:
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            if paragraph.runs:
                                paragraph.runs[0].font.size = Pt(10)

                    # Iterar sobre las filas de datos en el grupo
                    for row_data in filas:
                        tipo_fila = row_data[-1]  # Última columna que define el tipo (1 o 2)
                        row = table.add_row().cells  # Crear nueva fila

                        if tipo_fila == 1:
                            # Fila tipo 1: "SI/NO" con comentarios
                            descripcion = row_data[3]
                            condicion = row_data[4]
                            comentario = row_data[7]

                            row[0].text = descripcion
                            row[1].text = "SI"
                            row[2].text = "(X)" if condicion == "SI" else "()"
                            row[3].text = "NO"
                            row[4].text = "(X)" if condicion == "NO" else "()"
                            row[5].text = comentario

                            # Pintar celdas según la condición
                            if "(X)" in row[2].text:
                                ReporteAnexo1.set_cell_background_color(row[1], "C6E0B4")
                                ReporteAnexo1.set_cell_background_color(row[2], "C6E0B4")
                            if "(X)" in row[4].text:
                                ReporteAnexo1.set_cell_background_color(row[3], "C6E0B4")
                                ReporteAnexo1.set_cell_background_color(row[4], "C6E0B4")

                        elif tipo_fila == 2:
                            # Fila tipo 2: Descripción, Medidas de control, Plazo, Responsable
                            descripcion = row_data[3]
                            medidas_control = row_data[5]
                            plazo = row_data[6]
                            responsable = row_data[8]

                            row[0].text = descripcion
                            row[1].text = medidas_control
                            row[3].text = plazo
                            row[5].text = responsable

                            # Combinar celdas
                            row[1].merge(row[2])
                            row[3].merge(row[4])

                        # Alinear y ajustar tamaño de fuente en celdas
                        for j in range(6):
                            cell = row[j]
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            for paragraph in cell.paragraphs:
                                if j == 5:
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                else:
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if j in [1, 2, 3, 4] else WD_PARAGRAPH_ALIGNMENT.LEFT
                                if paragraph.runs:
                                    paragraph.runs[0].font.size = Pt(10)

                    # Centramos la tabla y agregamos espacio al final
                    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Separar las tablas con un espacio
                    doc.add_paragraph("")
    
    def generar_tablaFirma(doc, responsable, fecha_completa):
        # Crear la tabla con 5 filas y 1 columna
        table = doc.add_table(rows=5, cols=1)
        table.style = 'Table Grid'  # Estilo de tabla

        # Ajustar el ancho de la columna para que la imagen se ajuste correctamente
        table.columns[0].width = Inches(2.5)  # Ajustar el ancho de la columna

        # Insertar la imagen de la firma en la primera celda
        cell_firma = table.cell(0, 0)
        cell_firma.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Añadir un párrafo con un carácter invisible antes de la imagen para espaciado
        paragraph_before = cell_firma.add_paragraph()
        paragraph_before.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_before.space_before = Pt(0)  # Ajustar el espaciado antes del párrafo a 0 puntos
        paragraph_before.space_after = Pt(0)   # Ajustar el espaciado después del párrafo a 0 puntos
        run_before = paragraph_before.add_run('\u00A0')  # Carácter invisible
        run_before.font.size = Pt(0.5)  # Fuente de 1 punto

        paragraph = cell_firma.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.space_before = Pt(0)  # Ajustar el espaciado antes del párrafo a 0 puntos
        paragraph.space_after = Pt(0)   # Ajustar el espaciado después del párrafo a 0 puntos
        run = paragraph.add_run()

        # Convertir el blob en una imagen y agregarla al documento
        if responsable:
            blob_firma = responsable[6]
            if blob_firma:
                try:
                    image_stream = BytesIO(blob_firma)  # Convertir el blob a un objeto BytesIO
                    # Ajustar el tamaño de la imagen al ancho deseado
                    run.add_picture(image_stream, width=Inches(1.5))
                except ZeroDivisionError:
                    try:
                        img_stream = ReporteAnexo1.reparar_imagen(image_stream)
                        # Ajustar el tamaño de la imagen al ancho deseado
                        run.add_picture(img_stream, width=Inches(1.5))
                    except Exception as e:
                        paragraph.add_run(f"Error al insertar la imagen: {str(e)}")
                except Exception as e:
                    paragraph.add_run(f"Error al insertar la imagen: {str(e)}")

            data = [
                responsable[2],
                f"DNI {responsable[4]}",
                f"CIP {responsable[5]}",
                f"FECHA: {fecha_completa}"
            ]
        else:
            data = [
                "Nombres y Apellidos",
                "DNI 04055799",
                "CIP 75616",
                f"FECHA: {fecha_completa}"
            ]

        # Añadir un párrafo con un carácter invisible después de la imagen para espaciado
        paragraph_after = cell_firma.add_paragraph()
        paragraph_after.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_after.line_spacing = Pt(0)  # Ajustar el espaciado entre líneas a 0 puntos
        run_after = paragraph_after.add_run('\u00A0')  # Carácter invisible
        run_after.font.size = Pt(2)  # Fuente de 1 punto

        # Llenar las celdas con el texto
        for i, text in enumerate(data):
            cell = table.cell(i + 1, 0)
            cell.text = text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.line_spacing = Pt(0)  # Ajustar el espaciado entre líneas a 0 puntos
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Ajuste de tamaño de fuente
                # Mantener el párrafo junto
                ReporteAnexo1.set_keep_together(paragraph)

        # Alinear la tabla al borde derecho de la página
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_page_break()
        # Agregar párrafos vacíos para centrar verticalmente el título
        for _ in range(10):  # Ajustar el número de párrafos según sea necesario
            doc.add_paragraph()

        # Crear un párrafo centrado con el título "Misceláneos PLANOS"
        titulo_paragraph = doc.add_paragraph()
        titulo_run = titulo_paragraph.add_run("Misceláneos\nPLANOS")
        titulo_run.bold = True  # Aplicar negrita
        titulo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Ajustar el tamaño de fuente del título
        titulo_run.font.size = Pt(14)  # Ajuste de tamaño de fuente
        
    # Función para agregar el contenido de la quinta página con ajustes adicionales y alineación justificada
    def agregar_contenido_anexo1(componentes, anexo1, doc, responsable, fecha_completa):
        if anexo1:
            ReporteAnexo1.agregar_titulo_1(doc, f"ANEXO 1: {anexo1[19].upper()}", WD_PARAGRAPH_ALIGNMENT.LEFT)
            doc.add_paragraph()
            ReporteAnexo1.justificar_parrafo(doc, anexo1[2])
            ReporteAnexo1.agregar_viñeta_numerica_negrita(doc, '1. RESÚMEN EJECUTIVO')
            ReporteAnexo1.generar_tabla_resumen_ejecutivo(doc, componentes)
            ReporteAnexo1.agregar_viñeta_numerica_negrita(doc, '2. CONTROL DE PARÁMETROS DE DISEÑO (CONDICIONES GEOMÉTRICAS Y PARÁMETROS OPERATIVOS)')
            ReporteAnexo1.generar_tabla_parametros(doc, componentes)
            ReporteAnexo1.agregar_viñeta_numerica_negrita(doc,'3. VERIFICACIÓN DE CONDICIONES FÍSICAS')
            ReporteAnexo1.generar_tabla_condiciones_fisicas(doc, componentes)
            ReporteAnexo1.agregar_viñeta_numerica_negrita(doc,'4. OPERATIVIDAD DE EQUIPOS DE MONITOREO (INSTRUMENTACIÓN GEOTÉCNICA)')
            ReporteAnexo1.generar_tabla_operatividad_equipos(doc, componentes)
            ReporteAnexo1.agregar_viñeta_numerica_negrita(doc,'5. OBSERVACIONES, MEDIDAS ADOPTADAS Y SEGUIMIENTO')
            ReporteAnexo1.generar_tabla_observaciones(doc, componentes)
            ReporteAnexo1.generar_tablaFirma(doc, responsable, fecha_completa)
    
    # Configuración principal del documento
    def configurar_documento_apa(proyectoid, fecha_completa, mes_anio, dia_mes_anio):
        respuesta = ReporteController.ctrlListarDatosGeneralAnexos(proyectoid, 'Anexo1')
        componentes = ReporteController.ctrlObtenerComponentes(proyectoid)
        ReporteAnexo1.doc = ReporteAnexo1.crear_documento()
        ReporteAnexo1.configurar_fuente(ReporteAnexo1.doc)
        datos_generales=EmpresaController.ctrlObtenerDatosConfiguracionEmpresa()
        ReporteAnexo1.agregar_caratula(ReporteAnexo1.doc, datos_generales[5], respuesta, mes_anio)
        ReporteAnexo1.agregar_encabezado(ReporteAnexo1.doc, datos_generales)
        ReporteAnexo1.agregar_pie_pagina(ReporteAnexo1.doc)
        responsable = EmpresaController.ctrlObtenerDatosConfiguracionResponsable(proyectoid)
        ReporteAnexo1.agregar_contenido_memorandum(ReporteAnexo1.doc, respuesta, responsable, datos_generales[1], dia_mes_anio)
        ReporteAnexo1.agregar_tercera_pagina(ReporteAnexo1.doc, respuesta, mes_anio)
        ReporteAnexo1.agregar_indice_con_marcador(ReporteAnexo1.doc)
        ReporteAnexo1.agregar_contenido_quinta_pagina(ReporteAnexo1.doc, respuesta)
        ReporteAnexo1.agregar_contenido_anexo1(componentes, respuesta, ReporteAnexo1.doc, responsable, fecha_completa)
    
    def generarReporte(proyectoid, fechainicio, fechafinal):
        timezone = pytz.timezone("America/Lima")
        fecha_actual = datetime.now(timezone)
        mes_anio = MetodosGenerales.obtenerMesAnio(fechafinal)
        dia_mes_anio = MetodosGenerales.obtenerDiaMesAnio(fecha_actual)
        fecha_formateada = fecha_actual.strftime("%d/%m/%Y")
        # actualizar indice
        ReporteAnexo1.configurar_documento_apa(proyectoid, fecha_formateada, mes_anio, dia_mes_anio)
        doc_path = "modules/reportes/ANEXO1.docx"
        pdf_path = "modules/reportes/ANEXO1.pdf"
        ActualizarReporte.guardar_y_actualizar_indice_y_ajustar_tablas(doc_path,pdf_path,ReporteAnexo1.doc)
    
    def guardarInformacionGeneralAnexo1(main, idproyecto, imagencomponente, tipoanexo):
        if idproyecto:
            # PORTADA
            titulo_portada = main.findChild(QLineEdit, "input_titulo_portada_A1").text()
            subtitulo_portada = main.findChild(QLineEdit, "input_subtitulo_portada_A1").text()
            lugar_portada = main.findChild(QLineEdit, "input_lugar_portada_A1").text()
            autor_portada = main.findChild(QLineEdit, "input_autor_portada_A1").text()
            # DOCUMENTO
            tipo_documento = main.findChild(QLineEdit, "input_tipo_documento_A1").text()
            codigo_reporte = main.findChild(QLineEdit, "input_codigo_reporte_A1").text()
            destinatario = main.findChild(QPlainTextEdit, "input_destinatario_reporte_A1").toPlainText()
            remitente = main.findChild(QPlainTextEdit, "input_remitente_reporte_A1").toPlainText()
            asunto = main.findChild(QPlainTextEdit, "input_asunto_reporte_A1").toPlainText()
            descripcion = main.findChild(QPlainTextEdit, "input_descripcion_reporte_A1").toPlainText()
            # COMPONENTE
            tipo_reporte = main.findChild(QLineEdit, "input_tipo_reporte_A1").text()
            nombre_componente = main.findChild(QLineEdit, "input_componente_reporte_A1").text()
            # INTRODUCCIÓN
            objetivo_reporte = main.findChild(QPlainTextEdit, "input_objetivo_reporte_A1").toPlainText()
            finalidad_reporte = main.findChild(QPlainTextEdit, "input_finalidad_reporte_A1").toPlainText()
            ambito_reporte = main.findChild(QPlainTextEdit, "input_ambito_reporte_A1").toPlainText()
            detalles_reporte = main.findChild(QPlainTextEdit, "input_detalle_reporte_A1").toPlainText()
            titulo_anexo1 = main.findChild(QLineEdit, "input_titulo_anexo_A1").text()
            # Agrupar los datos en una lista
            datos = [titulo_portada, subtitulo_portada, lugar_portada, autor_portada, tipo_documento, codigo_reporte,
                     destinatario, remitente, asunto, descripcion, tipo_reporte, nombre_componente, imagencomponente,
                     objetivo_reporte, finalidad_reporte, ambito_reporte, detalles_reporte, titulo_anexo1, tipoanexo]
            respuesta = ReporteController.ctrlGuardarDataGeneralAnexos(datos, idproyecto, tipoanexo)
            return respuesta
    
    def cargarDataFormulariosAnexoGeneral(main, idproyecto):
        general = ReporteController.ctrlListarDatosGeneralAnexos(idproyecto, "Anexo1")
        if general:
            main.findChild(QLineEdit, "input_titulo_portada_A1").setText(general[2])
            main.findChild(QLineEdit, "input_subtitulo_portada_A1").setText(general[3])
            main.findChild(QLineEdit, "input_lugar_portada_A1").setText(general[4])
            main.findChild(QLineEdit, "input_autor_portada_A1").setText(general[5])
            # DOCUMENTO
            main.findChild(QLineEdit, "input_tipo_documento_A1").setText(general[6])
            main.findChild(QLineEdit, "input_codigo_reporte_A1").setText(general[7])
            main.findChild(QPlainTextEdit, "input_destinatario_reporte_A1").setPlainText(general[8])
            main.findChild(QPlainTextEdit, "input_remitente_reporte_A1").setPlainText(general[9])
            main.findChild(QPlainTextEdit, "input_asunto_reporte_A1").setPlainText(general[10])
            main.findChild(QPlainTextEdit, "input_descripcion_reporte_A1").setPlainText(general[11])
            # COMPONENTE
            main.findChild(QLineEdit, "input_tipo_reporte_A1").setText(general[12])
            main.findChild(QLineEdit, "input_componente_reporte_A1").setText(general[13])
            if general[14]:
                pixmap = MetodosGenerales.convertir_blob_a_pixmap(general[14])
                scaled_pixmap = pixmap.scaled(120, 120, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                main.findChild(QLabel, "lb_imagen_componente_A1").setPixmap(scaled_pixmap)
            # INTRODUCCIÓN
            main.findChild(QPlainTextEdit, "input_objetivo_reporte_A1").setPlainText(general[15])
            main.findChild(QPlainTextEdit, "input_finalidad_reporte_A1").setPlainText(general[16])
            main.findChild(QPlainTextEdit, "input_ambito_reporte_A1").setPlainText(general[17])
            main.findChild(QPlainTextEdit, "input_detalle_reporte_A1").setPlainText(general[18])
            main.findChild(QLineEdit, "input_titulo_anexo_A1").setText(general[19])
    
    def guardarInformacionResumenEjecutivo(main, idproyecto):
        if idproyecto:
            fecha_hora_actual = datetime.now()
            fechahora = fecha_hora_actual.strftime("%Y-%m-%d %H:%M:%S") 
            idcomponente = main.findChild(QComboBox, "cb_componentes_anexos").currentData()
            descripcion_anexo = main.findChild(QPlainTextEdit, "input_descripcion_general_anexo_A1").toPlainText()
            componente_encabezado_a1 = main.findChild(QLineEdit, "input_componente_encabezado_anexo_A1").text()
            valor_componente_encabezado_a1 = main.findChild(QLineEdit, "input_valor_componente_encabezado_anexo_A1").text()
            autorizacion_encabezado_a1 = main.findChild(QLineEdit, "input_autorizacion_encabezado_anexo_A1").text()
            valor_autorizacion_encabezado_a1 = main.findChild(QLineEdit, "input_valor_autorizacion_encabezado_anexo_A1").text()
            fecha_encabezado_a1 = main.findChild(QLineEdit, "input_fecha_encabezado_anexo_A1").text()
            valor_fecha_encabezado_a1 = main.findChild(QLineEdit, "input_valor_fecha_encabezado_anexo_A1").text()
            expediente_control_a1 = main.findChild(QPlainTextEdit, "input_expediente_control_anexo_A1").toPlainText()
            si_expediente_a1 = main.findChild(QRadioButton, "rb_expediente_control_SI_anexo_A1")
            valor_expediente_a1 = 'SI' if si_expediente_a1.isChecked() else 'NO'
            inspeccion_a1 = main.findChild(QPlainTextEdit, "input_inspeccion_anexo_A1").toPlainText()
            si_inspeccion_a1 = main.findChild(QRadioButton, "rb_inspeccion_SI_anexo_A1")
            valor_inspeccion_a1 = 'SI' if si_inspeccion_a1.isChecked() else 'NO'
            # Enviar los datos al controlador
            valores = [descripcion_anexo, componente_encabezado_a1, valor_componente_encabezado_a1,
                autorizacion_encabezado_a1, valor_autorizacion_encabezado_a1, fecha_encabezado_a1, valor_fecha_encabezado_a1,
                expediente_control_a1, valor_expediente_a1, inspeccion_a1, valor_inspeccion_a1]
            respuesta = ReporteController.ctrlGuardarResumenEjecutivoAnexo1(valores, idcomponente)
            return respuesta
    
    def cargarInformacionResumenEjecutivoAnexo(main):
        idcomponente = main.findChild(QComboBox, "cb_componentes_anexos").currentData()
        resumen = ReporteController.ctrlObtenerResumenEjecutivoAnexo1(idcomponente)
        if resumen:
            main.findChild(QPlainTextEdit, "input_descripcion_general_anexo_A1").setPlainText(resumen[2])
            main.findChild(QLineEdit, "input_componente_encabezado_anexo_A1").setText(resumen[3])
            main.findChild(QLineEdit, "input_valor_componente_encabezado_anexo_A1").setText(resumen[4])
            main.findChild(QLineEdit, "input_autorizacion_encabezado_anexo_A1").setText(resumen[5])
            main.findChild(QLineEdit, "input_valor_autorizacion_encabezado_anexo_A1").setText(resumen[6])
            main.findChild(QLineEdit, "input_fecha_encabezado_anexo_A1").setText(resumen[7])
            main.findChild(QLineEdit, "input_valor_fecha_encabezado_anexo_A1").setText(resumen[8])
            main.findChild(QPlainTextEdit, "input_expediente_control_anexo_A1").setPlainText(resumen[9])
            si_expediente_a1 = main.findChild(QRadioButton, "rb_expediente_control_SI_anexo_A1")
            no_expediente_a1 = main.findChild(QRadioButton, "rb_expediente_control_NO_anexo_A1")
            if resumen[10] == 'SI':
                si_expediente_a1.setChecked(True)
            else:
                no_expediente_a1.setChecked(True)
            main.findChild(QPlainTextEdit, "input_inspeccion_anexo_A1").setPlainText(resumen[11])
            si_inspeccion_a1 = main.findChild(QRadioButton, "rb_inspeccion_SI_anexo_A1")
            no_inspeccion_a1 = main.findChild(QRadioButton, "rb_inspeccion_NO_anexo_A1")
            if resumen[12] == 'SI':
                si_inspeccion_a1.setChecked(True)
            else:
                no_inspeccion_a1.setChecked(True)
    
    def guardarDatosDinamicosAnexo1(widget_anexo1, componente):
        parametros_a1 = ReporteAnexo1.obtener_valores_frame_parametros_A1(widget_anexo1)
        if parametros_a1:
            ReporteController.ctrlGuardarParametrosAnexo1(componente, parametros_a1)
        condiciones_a1 = ReporteAnexo1.obtener_valores_frame_condiciones_A1(widget_anexo1)
        if condiciones_a1:
            ReporteController.ctrlGuardarCondicionesAnexo1(componente, condiciones_a1)
        operatividad_a1 = ReporteAnexo1.obtener_valores_frame_operatividad_A1(widget_anexo1)
        if operatividad_a1:
            ReporteController.ctrlGuardarOperatividadAnexo1(componente, operatividad_a1)
        observaciones_a1 = ReporteAnexo1.obtener_valores_frame_observaciones_A1(widget_anexo1)
        if observaciones_a1:
            ReporteController.ctrlGuardarObservacionesAnexo1(componente, observaciones_a1)
    
    def obtener_valores_frame_parametros_A1(widget):
        # Buscar el layout principal del widget
        layout = widget.layout()
        # Acceder al frame3 (que está dentro del layout, en el índice 2)
        frame = layout.itemAt(0).widget()  # Obtener el frame3 (índice 2)
        # Obtener el layout dentro de frame3
        frame_layout = frame.layout()
        # Lista para almacenar los valores de las filas
        valores = []
        # Recorremos las filas, comenzando desde la fila 2 (porque la fila 1 tiene los encabezados)
        for i in range(2, frame_layout.rowCount()):
            row_values = []
            # Celda 0: QLineEdit (DESCRIPCIÓN)
            line_edit_desc = frame_layout.itemAtPosition(i, 0).widget()
            descripcion = line_edit_desc.text() if line_edit_desc else ""
            # Verificamos si la descripción está vacía; si es así, no consideramos la fila
            if not descripcion:
                continue  # Saltamos esta iteración si la descripción está vacía
            row_values.append(descripcion)
            # Celda 1: QLineEdit (PARÁMETRO 1)
            line_edit_param1 = frame_layout.itemAtPosition(i, 1).widget()
            row_values.append(line_edit_param1.text() if line_edit_param1 else "")
            # Celda 2: QLineEdit (PARÁMETRO 2)
            line_edit_param2 = frame_layout.itemAtPosition(i, 2).widget()
            row_values.append(line_edit_param2.text() if line_edit_param2 else "")
            # Celda 3 y 4: Buscar los QRadioButton en el QGroupBox (CONDICIÓN ACTUAL)
            group_box_condition = frame_layout.itemAtPosition(i, 3).widget()
            # Buscar todos los widgets dentro del group box
            radio_buttons = [widget for widget in group_box_condition.findChildren(QRadioButton)]
            # Verificar cuál radio button está seleccionado
            if radio_buttons:
                if any(radio.isChecked() for radio in radio_buttons):
                    # Si alguno está seleccionado, agregar el valor correspondiente
                    selected_radio = next((radio.text() for radio in radio_buttons if radio.isChecked()), "")
                    row_values.append(selected_radio)
                else:
                    row_values.append("")  # Si ninguno está seleccionado, dejamos el valor vacío
            else:
                row_values.append("")  # Si no hay radio buttons, agregamos un valor vacío
            # Celda 5: QPlainTextEdit (COMENTARIOS)
            text_edit_comments = frame_layout.itemAtPosition(i, 5).widget()
            row_values.append(text_edit_comments.toPlainText() if text_edit_comments else "")
            # Añadir los valores de esta fila a la lista de valores
            valores.append(row_values)
        return valores
    
    def obtener_valores_frame_condiciones_A1(widget):
        # Buscar el layout principal del widget
        layout = widget.layout()
        # Acceder al frame4 (que está dentro del layout, en el índice 3)
        frame = layout.itemAt(1).widget()  # Obtener el frame4 (índice 3)
        # Obtener el layout dentro de frame4
        frame_layout = frame.layout()
        # Lista para almacenar los valores de las filas
        valores = []
        # Recorremos las filas, comenzando desde la fila 2 (porque las filas 0 y 1 son encabezados)
        for i in range(2, frame_layout.rowCount()):
            row_values = []
            # Celda 0: QLineEdit (Condición)
            line_edit_condition = frame_layout.itemAtPosition(i, 0).widget()
            if not line_edit_condition or not line_edit_condition.text().strip():
                # Si la celda está vacía, omitimos esta fila
                continue
            row_values.append(line_edit_condition.text().strip())  # Agregamos el texto del campo Condición
            # Celda 1: Condición Actual (puede ser QGroupBox o QLineEdit dependiendo del tipo de fila)
            condition_widget = frame_layout.itemAtPosition(i, 1).widget()
            tipo = None
            if isinstance(condition_widget, QGroupBox):  # Si es un QGroupBox (con radio buttons)
                # Buscar todos los QRadioButton dentro del group box
                radio_buttons = condition_widget.findChildren(QRadioButton)
                # Verificar cuál está seleccionado
                if any(radio.isChecked() for radio in radio_buttons):
                    selected_radio = next((radio.text() for radio in radio_buttons if radio.isChecked()), "")
                    row_values.append(selected_radio)
                else:
                    row_values.append("")  # Si ninguno está seleccionado, dejamos vacío
                # Determinar tipo según los textos de los radio buttons
                radio_texts = [radio.text().strip().upper() for radio in radio_buttons]
                if "CUMPLE" in radio_texts and "NO CUMPLE" in radio_texts:
                    tipo = 1
                elif "SI" in radio_texts or "SI" in radio_texts and "NO" in radio_texts:
                    tipo = 2
            elif isinstance(condition_widget, QLineEdit):  # Si es un QLineEdit
                row_values.append(condition_widget.text().strip() if condition_widget else "")
                tipo = 3  # Tipo 3: solo inputs de texto
            # Celda 2: QPlainTextEdit (Comentarios)
            text_edit_comments = frame_layout.itemAtPosition(i, 2).widget()
            row_values.append(text_edit_comments.toPlainText().strip() if text_edit_comments else "")
            # Añadir tipo al final de la fila
            row_values.append(tipo)
            # Añadir los valores de esta fila a la lista de valores
            valores.append(row_values)
        return valores

    def obtener_valores_frame_operatividad_A1(widget):
        # Buscar el layout principal del widget
        layout = widget.layout()
        # Acceder al frame5 (que está dentro del layout, en el índice 4)
        frame = layout.itemAt(2).widget()  # Obtener el frame5 (índice 4)
        # Obtener el layout dentro de frame5
        frame_layout = frame.layout()
        # Lista para almacenar los valores de las filas
        valores = []
        # Recorremos las filas, comenzando desde la fila 3 (porque las filas 0, 1 y 2 son encabezados)
        for i in range(3, frame_layout.rowCount()):
            row_values = []
            # Celda 0: QLineEdit (Instrumentación)
            line_edit_instrumentation = frame_layout.itemAtPosition(i, 0).widget()
            if not line_edit_instrumentation or not line_edit_instrumentation.text().strip():
                # Si la celda está vacía, omitimos esta fila
                continue
            row_values.append(line_edit_instrumentation.text().strip())  # Agregar Instrumentación
            # Celda 1: Condición Actual (QGroupBox con radio buttons)
            group_box_condition = frame_layout.itemAtPosition(i, 1).widget()
            radio_buttons = group_box_condition.findChildren(QRadioButton)
            if any(radio.isChecked() for radio in radio_buttons):
                selected_radio = next((radio.text() for radio in radio_buttons if radio.isChecked()), "")
                row_values.append(selected_radio)
            else:
                row_values.append("")  # Si ninguno está seleccionado, dejamos vacío
            # Celda 2: QSpinBox (Cantidad)
            spin_box_quantity = frame_layout.itemAtPosition(i, 2).widget()
            row_values.append(spin_box_quantity.value() if spin_box_quantity else 0)
            # Celda 3: QCheckBox (Operativo)
            checkbox_operativo = frame_layout.itemAtPosition(i, 3).widget()
            row_values.append("SI" if checkbox_operativo.isChecked() else "")  # Marcar como "SI" si está checked, "" si no
            # Celda 4: QLineEdit (Comentarios)
            line_edit_comments = frame_layout.itemAtPosition(i, 4).widget()
            row_values.append(line_edit_comments.text().strip() if line_edit_comments else "")
            # Añadir los valores de esta fila a la lista de valores
            valores.append(row_values)
        return valores
    
    def obtener_valores_frame_observaciones_A1(widget):
        # Buscar el layout principal del widget
        layout = widget.layout()
        frame = layout.itemAt(3).widget() 
        # Obtener el layout dentro de frame6
        frame_layout = frame.layout()
        # Lista para almacenar los valores de las filas
        valores = []
        # Recorremos las filas dinámicas, comenzando desde la fila 2 (porque las filas 0 y 1 son encabezados)
        for i in range(2, frame_layout.rowCount()):
            row_values = []
            # Celda 0: QPlainTextEdit (Descripción)
            description_widget = frame_layout.itemAtPosition(i, 0).widget()
            if not description_widget or not description_widget.toPlainText().strip():
                # Si la descripción está vacía, omitir esta fila
                continue
            descripcion = description_widget.toPlainText().strip()
            row_values.append(descripcion)  # Agregar Descripción
            # Determinar el tipo de fila (SI/NO o Comentarios Extendidos)
            condition_widget = frame_layout.itemAtPosition(i, 1).widget()
            if isinstance(condition_widget, QGroupBox):
                # Es una fila SI/NO
                radio_buttons = condition_widget.findChildren(QRadioButton)
                condicion_actual = next((radio.text() for radio in radio_buttons if radio.isChecked()), "")
                row_values.append(condicion_actual)  # Condición Actual
                row_values.append("")  # Medidas de Control vacío
                row_values.append("")  # Plazo vacío
                comments_widget = frame_layout.itemAtPosition(i, 3).widget()
                comentario = comments_widget.toPlainText().strip() if comments_widget else ""
                row_values.append(comentario)  # Comentarios
                row_values.append("")  # Responsable vacío
                row_values.append(1)  # Tipo de fila: 1 (SI/NO)
            else:
                # Es una fila de Comentarios Extendidos
                row_values.append("")  # Condición Actual vacío
                medidas_control_widget = frame_layout.itemAtPosition(i, 1).widget()
                medidas_control = medidas_control_widget.toPlainText().strip() if medidas_control_widget else ""
                row_values.append(medidas_control)  # Medidas de Control
                plazo_widget = frame_layout.itemAtPosition(i, 2).widget()
                plazo = plazo_widget.toPlainText().strip() if plazo_widget else ""
                row_values.append(plazo)  # Plazo
                comments_widget = frame_layout.itemAtPosition(i, 3).widget()
                responsable = comments_widget.toPlainText().strip() if comments_widget else ""
                row_values.append("")  # Comentarios vacío
                row_values.append(responsable)  # Responsable
                row_values.append(2)  # Tipo de fila: 2 (Comentarios Extendidos)
            # Añadir los valores de esta fila a la lista de valores
            valores.append(row_values)
        return valores
    
    def mostrar_mensaje(titulo, mensaje, icono):
        msg_box = QMessageBox()
        msg_box.setWindowTitle(titulo)
        msg_box.setText(mensaje)
        msg_box.setIcon(icono)
        msg_box.exec()
    