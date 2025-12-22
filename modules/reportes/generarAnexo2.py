import pandas as pd
import pytz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from io import BytesIO
from PySide6.QtWidgets import (QMessageBox, QLineEdit, QPlainTextEdit, QLabel, QComboBox, QRadioButton, QGroupBox)
from datetime import datetime
from PIL import Image
from PySide6.QtCore import Qt
from utils.common.metodosGenerales import MetodosGenerales
from controllers.ReporteController import ReporteController
from controllers.UmbralController import UmbralController
from controllers.InclinometroController import InclinometroController
from controllers.PiezometroController import PiezometroController
from controllers.EmpresaController import EmpresaController
from controllers.InstrumentacionController import InstrumentacionController
from controllers.CeldaController import CeldaController
from controllers.TDRController import TDRController
from modules.empresa.softwareconfiguracion import SoftwareConfiguracion
from utils.shared.actualizarReporte import ActualizarReporte
class ReporteAnexo2:
    doc = None
    
    # Función que justifica el párrafo en base a la longitud en caracteres
    def justificar_parrafo(doc, texto, max_caracteres_por_linea=80):
        # Separar el texto por saltos de línea
        lineas = texto.split('\n')
        for idx, linea in enumerate(lineas):
            parrafo = doc.add_paragraph(linea)
            # Calcular la longitud en caracteres de la línea
            longitud_linea = len(linea)
            # Justificar si la línea excede la longitud máxima en caracteres
            if longitud_linea > max_caracteres_por_linea:
                parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            else:
                parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    def agregar_parrafos_con_vinetas(doc, texto):
        # Dividir el texto en párrafos usando el salto de línea
        parrafos = texto.split('\n')
        for parrafo in parrafos:
            # Agregar cada párrafo con una viñeta
            nuevo_parrafo = doc.add_paragraph(parrafo, style='List Bullet')
            # Ajustar la sangría y alineación
            nuevo_parrafo.paragraph_format.left_indent = Inches(0.5)
            nuevo_parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    def agregar_parrafos_alternando_negrita(cell, texto):
        # Dividir el texto en párrafos usando el salto de línea
        parrafos = texto.split('\n')
        # Eliminar cualquier párrafo vacío al principio
        parrafos = [p for p in parrafos if p.strip()]
        # Obtener el primer párrafo existente en la celda (si existe)
        if len(cell.paragraphs) > 0:
            nuevo_parrafo = cell.paragraphs[0]  # Usar el primer párrafo existente
        else:
            # Si no hay párrafos existentes, crear uno nuevo
            nuevo_parrafo = cell.add_paragraph()
        # Recorrer los párrafos divididos y agregar los textos con alternancia de estilos
        for i, parrafo in enumerate(parrafos):
            # Si no es el primer párrafo, agregar un salto de línea
            if i > 0:
                nuevo_parrafo.add_run("\n")
            # Eliminar los espacios innecesarios antes de agregar el texto
            run = nuevo_parrafo.add_run(parrafo.strip())
            # Alternar entre negrita y normal
            if i % 2 == 0:  # Para párrafos con índice par (negrita)
                run.bold = True
            else:  # Para párrafos con índice impar (normal)
                run.bold = False
    
    def set_cell_background_color(cell, color_hex):
        """Aplica un color de fondo en hexadecimal a una celda de la tabla."""
        cell_properties = cell._element.get_or_add_tcPr()
        cell_shading = OxmlElement('w:shd')
        cell_shading.set(qn('w:fill'), color_hex)  # Configura el color de fondo usando qn para el espacio de nombres
        cell_properties.append(cell_shading)
    
    def agregar_viñeta_numerica_negrita(doc, texto, estilo="Heading 2"):
        # Crear el párrafo con el estilo indicado
        parrafo = doc.add_paragraph(texto, style=estilo)
        # Aplicar formato al texto (negrita y tamaño de fuente)
        for run in parrafo.runs:
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
        # Alineación (opcional)
        parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # Configurar el párrafo para mantenerse junto con la siguiente tabla
        ReporteAnexo2.set_keep_together(parrafo)
        return parrafo
    
    def agregar_viñeta_letras_negrita(doc, texto, letra):
        # Crear un párrafo
        parrafo = doc.add_paragraph()

        # Agregar la viñeta en letra con formato en negrita
        run_letra = parrafo.add_run(f"{letra}. ")
        run_letra.font.bold = True  # Negrita
        run_letra.font.size = Pt(10)  # Tamaño de fuente ajustable

        # Agregar el texto principal en negrita
        run_texto = parrafo.add_run(texto)
        run_texto.font.bold = True  # Negrita
        run_texto.font.size = Pt(10)  # Tamaño de fuente ajustable

        # Alineación del párrafo
        parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Alineado a la izquierda
        
        # Configurar el párrafo para mantenerse junto con la siguiente tabla
        ReporteAnexo2.set_keep_together(parrafo)
    
    # para que siempre este juntos algunos elementos
    def set_keep_together(parrafo):
        """
        Configura el párrafo para mantenerse junto con el siguiente elemento (tabla, párrafo, etc.).
        """
        p = parrafo._p  # Accede al elemento XML subyacente del párrafo
        pPr = p.get_or_add_pPr()
        keepNext = OxmlElement('w:keepNext')  # Mantener con el siguiente párrafo o elemento
        pPr.append(keepNext)
    
    def insertar_imagen_con_descripcion(doc, imagen_blob, descripcion, numero_figura):
        # Añadir la imagen centrada
        par_imagen = doc.add_paragraph()
        par_imagen.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_imagen = par_imagen.add_run()
        run_imagen.add_picture(imagen_blob, width=Inches(4))  # Ajustar tamaño según necesidad

        # Añadir la descripción centrada en negrita
        texto_figura = f"Figura N° {numero_figura}: {descripcion}"
        par_descripcion = doc.add_paragraph(texto_figura)
        par_descripcion.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_descripcion = par_descripcion.runs[0]
        run_descripcion.bold = True
        run_descripcion.font.size = Pt(10)  # Tamaño de letra ajustable
    
    def reparar_imagen(img_stream):
        try:
            # Abrir la imagen usando PIL
            img = Image.open(img_stream)

            # Verificar las dimensiones de la imagen
            width, height = img.size
            if width == 0 or height == 0:
                raise ValueError("Dimensiones de la imagen no válidas")

            # Crear un nuevo BytesIO para guardar la imagen reparada
            repaired_img_stream = BytesIO()
            img.save(repaired_img_stream, format=img.format)
            repaired_img_stream.seek(0)

            return repaired_img_stream
        except Exception as e:
            raise ValueError(f"Error al reparar la imagen: {e}")

    # Crear un documento nuevo
    def crear_documento():
        return Document()

    # Configurar página en tamaño A4 y aplicar márgenes APA
    def configurar_pagina(section):
        section.page_height = Inches(11.69)  # Tamaño A4 en pulgadas (29.7 cm)
        section.page_width = Inches(8.27)    # Tamaño A4 en pulgadas (21 cm)
        section.top_margin = Inches(1.0)     # Margen superior APA (1 pulgada)
        section.bottom_margin = Inches(1.0)  # Margen inferior APA (1 pulgada)
        section.left_margin = Inches(1.0)    # Margen izquierdo APA (1 pulgada)
        section.right_margin = Inches(1.0)   # Margen derecho APA (1 pulgada)

    # Configurar estilo de fuente y tamaño según APA
    def configurar_fuente(doc):
        style = doc.styles['Normal']
        style.font.name = 'Arial'  # Fuente APA estándar
        style.font.size = Pt(12) # Tamaño de fuente APA
        
    # Función para crear un título con formato Heading 1 y espaciado automático
    def agregar_titulo_1(doc, texto, alineacion=WD_PARAGRAPH_ALIGNMENT.CENTER):
        # Agregar un párrafo vacío antes del título para espacio adicional

        # Crear el párrafo del título y aplicar el estilo y formato
        titulo = doc.add_paragraph(texto)
        titulo.style = 'Heading 1'
        titulo.alignment = alineacion  # Aplicar la alineación directamente

        # Aplicar formato específico: Arial, tamaño, color, negrita
        run = titulo.runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(12)  # Tamaño APA, ajusta según la norma
        run.font.color.rgb = RGBColor(0, 0, 0)  # Color negro
        run.bold = True  # Negrita si es requerido

        # Ajustar espacios antes y después del título
        titulo.paragraph_format.space_before = Pt(0)  # Sin espacio extra antes
        titulo.paragraph_format.space_after = Pt(0)  # Espacio después del título

        return titulo

    # Función para crear un título con formato Heading 2
    def agregar_titulo_2(doc, texto, alineacion=WD_PARAGRAPH_ALIGNMENT.LEFT):
        # Crear el párrafo del título y aplicar el estilo y formato
        titulo = doc.add_paragraph(texto)
        titulo.style = 'Heading 2'
        titulo.alignment = alineacion  # Aplicar la alineación directamente

        # Aplicar formato específico: Arial, tamaño, color, sin ajuste de espacios
        run = titulo.runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(12)  # Tamaño para título 2, ajusta si es necesario
        run.font.color.rgb = RGBColor(0, 0, 0)  # Color negro
        run.bold = True  # Negrita si es requerido

        return titulo
    
    ###############################################################################33
    # Añadir encabezado en la segunda sección del documento
    def agregar_encabezado(doc, datos_generales):
        if datos_generales:
            logotipo_blob=datos_generales[5]
            section = doc.sections[1]
            header = section.header
            header.is_linked_to_previous = False
            header_table = header.add_table(rows=1, cols=2, width=Inches(7))
            header_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for cell in header_table.rows[0].cells:
                cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            cell_logo = header_table.cell(0, 0)
            if logotipo_blob:
                try:
                    # Convertir el blob a un archivo en memoria y agregarlo como imagen
                    image_stream = BytesIO(logotipo_blob)
                    run = cell_logo.paragraphs[0].add_run()
                    # Ajustar el tamaño de la imagen al ancho deseado
                    run.add_picture(image_stream, width=Inches(1.5))
                except ZeroDivisionError:
                    try:
                        img_stream = ReporteAnexo2.reparar_imagen(image_stream)
                        # Ajustar el tamaño de la imagen al ancho deseado
                        run.add_picture(img_stream, width=Inches(1.5))
                    except Exception as e:
                        cell_logo.text = f"Error al cargar el logotipo: {str(e)}"
                except Exception as e:
                    cell_logo.text = f"Error al cargar el logotipo: {str(e)}"
                        
            else:
                cell_logo.text = " "

            cell_text = header_table.cell(0, 1)
            header_text = cell_text.paragraphs[0]
            header_text.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            run_text = header_text.add_run(f"{datos_generales[1]}\nJefatura de Geotecnia")
            run_text.font.size = Pt(10)
            run_text.bold = True

            tbl_pr = header_table._element.find(qn("w:tblPr"))
            if tbl_pr is None:
                tbl_pr = OxmlElement("w:tblPr")
                header_table._element.insert(0, tbl_pr)
            tbl_borders = OxmlElement("w:tblBorders")
            bottom_border = OxmlElement("w:bottom")
            bottom_border.set(qn("w:val"), "single")
            bottom_border.set(qn("w:sz"), "15")
            bottom_border.set(qn("w:space"), "0")
            bottom_border.set(qn("w:color"), "263656")
            tbl_borders.append(bottom_border)
            tbl_pr.append(tbl_borders)

            # Agregar un párrafo vacío debajo de la tabla
            header.add_paragraph("")
        
    # Configurar pie de página con borde superior y número de página en la segunda sección
    def agregar_pie_pagina(doc):
        section = doc.sections[1]
        section.footer.is_linked_to_previous = False

        footer = section.footer
        footer_table = footer.add_table(rows=1, cols=1, width=Inches(7))
        footer_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_cell = footer_table.cell(0, 0)

        page_number_paragraph = footer_cell.paragraphs[0]
        page_number_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        page_num_run = page_number_paragraph.add_run()
        page_num_run.font.size = Pt(10)

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        page_num_run._r.append(fldChar)

        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE \\* Arabic \\* MERGEFORMAT"
        page_num_run._r.append(instrText)

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "end")
        page_num_run._r.append(fldChar)

        tbl_pr_footer = footer_table._element.find(qn("w:tblPr"))
        if tbl_pr_footer is None:
            tbl_pr_footer = OxmlElement("w:tblPr")
            footer_table._element.insert(0, tbl_pr_footer)
        tbl_borders_footer = OxmlElement("w:tblBorders")
        top_border = OxmlElement("w:top")
        top_border.set(qn("w:val"), "single")
        top_border.set(qn("w:sz"), "15")
        top_border.set(qn("w:space"), "0")
        top_border.set(qn("w:color"), "263656")
        tbl_borders_footer.append(top_border)
        tbl_pr_footer.append(tbl_borders_footer)
    
    # Función para agregar caratula sin pie ni encabezado
    def agregar_caratula(doc, logotipo_blob, datos, mes_anio):
        # Acceder a la primera sección del documento
        if datos:
            section = doc.sections[0]
            ReporteAnexo2.configurar_pagina(section)
            section.different_first_page_header_footer = True  # Sin encabezado ni pie en la primera página

            # Añadir espacio vacío
            for _ in range(5):
                doc.add_paragraph("")

            # Si se pasa un BLOB de la imagen, usar BytesIO para insertarlo
            if logotipo_blob:
                try:
                    # Crear un objeto BytesIO a partir del BLOB
                    image_stream  = BytesIO(logotipo_blob)

                    # Intentar insertar la imagen en la carátula
                    img_paragraph = doc.add_paragraph()
                    img_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = img_paragraph.add_run()
                    run.add_picture(image_stream , width=Inches(2.5))
                except ZeroDivisionError:
                    # Si hay un error de división por cero, intentar reparar la imagen
                    try:
                        img_stream = ReporteAnexo2.reparar_imagen(image_stream)
                        run.add_picture(img_stream, width=Inches(2.5))
                    except Exception as e:
                        doc.add_paragraph(f"Error al insertar la imagen: {e}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                except Exception as e:
                    doc.add_paragraph(f"Error al insertar la imagen: {e}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                doc.add_paragraph("No se proporcionó imagen.").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Título del reporte
            title = doc.add_paragraph(datos[2])
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title.runs[0].bold = True
            title.runs[0].font.size = Pt(20)
            title.add_run("\n")

            # Subtítulo del reporte
            subtitle = doc.add_paragraph(f"{datos[3]}\nMES: {mes_anio}")
            subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            subtitle.runs[0].font.size = Pt(18)

            # Añadir más espacio vacío
            for _ in range(6):
                doc.add_paragraph("")

            # Pie de página
            footer_paragraph = doc.add_paragraph(f"{datos[4]}\nElaborado por: {datos[5]}")
            footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            footer_paragraph.runs[0].bold = True
            footer_paragraph.runs[0].font.size = Pt(11)

            # Crear una nueva sección para la siguiente página
            new_section = doc.add_section(WD_SECTION.NEW_PAGE)
            ReporteAnexo2.configurar_pagina(new_section)
            new_section.different_first_page_header_footer = False  # Permitir encabezado y pie
    
    # Añadir el contenido del "MEMORANDUM" en la segunda página utilizando tabla para alinear texto
    def agregar_contenido_memorandum(doc, datos, responsable, unidad_minera, dia_mes_anio):
        if datos:
            invisible_space = "\u00A0" * 10
            firma_blob = None
            if responsable:
                firma_blob = responsable[6]
            # Título del Docuemnto
            titulo = doc.add_paragraph(datos[6])
            titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            titulo_run = titulo.runs[0]
            titulo_run.font.size = Pt(14)
            titulo_run.bold = True

            # CODIGO
            referencia = doc.add_paragraph(datos[7])
            referencia.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            referencia.runs[0].font.size = Pt(10)
            referencia.runs[0].font.bold = True 

            # Tabla invisible para alinear texto de forma organizada
            table = doc.add_table(rows=4, cols=2)
            table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            table.autofit = False
            table.columns[0].width = Inches(1.5)
            table.columns[1].width = Inches(5.0)

            # Reducir los márgenes internos de las celdas para minimizar el espacio entre columnas
            for row in table.rows:
                for cell in row.cells:
                    tc_pr = cell._element.get_or_add_tcPr()
                    tc_mar = OxmlElement('w:tcMar')
                    for side in ['top', 'left', 'bottom', 'right']:
                        margin = OxmlElement(f'w:{side}')
                        margin.set(qn('w:w'), '0')  # Sin margen
                        margin.set(qn('w:type'), 'dxa')
                        tc_mar.append(margin)
                    tc_pr.append(tc_mar)

            # Rellenar la tabla con texto ajustado
            cell_0_0 = table.cell(0, 0).paragraphs[0].add_run(f"Para:{invisible_space}")
            cell_0_0.bold = True
            # cell_0_1 = table.cell(0, 1).paragraphs[0]
            # cell_0_1.add_run(datos[9])
            # Añadir texto alternando negrita y normal a la segunda celda de la fila
            cell_0_1 = table.cell(0, 1)
            ReporteAnexo2.agregar_parrafos_alternando_negrita(cell_0_1, datos[8])

            cell_1_0 = table.cell(1, 0).paragraphs[0].add_run("De: ")
            cell_1_0.bold = True
            cell_1_1 = table.cell(1, 1)
            ReporteAnexo2.agregar_parrafos_alternando_negrita(cell_1_1, datos[9])

            cell_2_0 = table.cell(2, 0).paragraphs[0].add_run("Fecha: ")
            cell_2_0.bold = True
            cell_2_1 = table.cell(2, 1).paragraphs[0]
            cell_2_1.add_run(dia_mes_anio)

            cell_3_0 = table.cell(3, 0).paragraphs[0].add_run("Asunto: ")
            cell_3_0.bold = True
            cell_3_1 = table.cell(3, 1).paragraphs[0]
            cell_3_1.add_run(datos[10])

            # Añadir el cuerpo del memorandum
            ReporteAnexo2.justificar_parrafo(doc,datos[11])

            # Añadir espacio antes de la firma
            for _ in range(2):
                doc.add_paragraph("")

            # Añadir la firma y detalles
            if firma_blob:                
                try:
                    # Convertir el blob a un archivo en memoria y agregarlo como imagen
                    firma_stream = BytesIO(firma_blob)
                    # Insertar la firma en el documento
                    firma_paragraph = doc.add_paragraph()
                    firma_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    firma_run = firma_paragraph.add_run()
                    firma_run.add_picture(firma_stream, width=Inches(1.5))
                except ZeroDivisionError:
                    try:
                        fir_stream = ReporteAnexo2.reparar_imagen(firma_stream)
                        # Insertar la firma en el documento
                        firma_paragraph = doc.add_paragraph()
                        firma_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        firma_run = firma_paragraph.add_run()
                        firma_run.add_picture(fir_stream, width=Inches(1.5))
                    except Exception as e:
                        doc.add_paragraph(f"Error al insertar la imagen: {e}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                except Exception as e:
                    doc.add_paragraph(f"Error al insertar la firma: {e}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                doc.add_paragraph("Firma no proporcionada.").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            # Información adicional sobre la firma
            if responsable:
                firma_paragraph = doc.add_paragraph(f"{responsable[2]}\n{responsable[3]}\n{unidad_minera}")
                firma_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            # Insertar salto de página
            doc.add_page_break()
        
    # Función para agregar la tercera página con la imagen y los títulos
    def agregar_tercera_pagina(doc, respuesta, mes_anio):
        if respuesta:
            # tipo reporte
            titulo = doc.add_paragraph(respuesta[12])
            titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            titulo.runs[0].bold = True
            titulo.runs[0].font.size = Pt(14)
            # fecha
            subtitulo = doc.add_paragraph(mes_anio)
            subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            subtitulo.runs[0].bold = True
            subtitulo.runs[0].font.size = Pt(12)
            # componente
            componente = doc.add_paragraph(respuesta[13])
            componente.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            componente.runs[0].bold = True
            componente.runs[0].font.size = Pt(12)
            doc.add_paragraph("")
            # Insertar la imagen BLOB en el centro de la página
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            MAX_WIDTH = 6.0  # pulgadas
            MAX_HEIGHT = 7.0  # pulgadas
            if respuesta[14]:
                try:
                    image_stream = BytesIO(respuesta[14])
                    # Abrir imagen con PIL para obtener dimensiones
                    image = Image.open(image_stream)
                    width_px, height_px = image.size
                    dpi = image.info.get('dpi', (96, 96))[0]  # usar 96 dpi por defecto
                    width_in = width_px / dpi
                    height_in = height_px / dpi
                    # Escalar proporcionalmente para ajustarse a los límites máximos
                    scale = min(MAX_WIDTH / width_in, MAX_HEIGHT / height_in)
                    final_width = Inches(width_in * scale)
                    final_height = Inches(height_in * scale)
                    # Reiniciar el stream para Word (PIL lo ha movido)
                    image_stream.seek(0)
                    run.add_picture(image_stream, width=final_width, height=final_height)
                except ZeroDivisionError:
                    try:
                        img_stream = ReporteAnexo2.reparar_imagen(image_stream)
                        run.add_picture(img_stream, width=Inches(6.5))
                    except Exception as e:
                        paragraph.add_run(f"Error al insertar la imagen: {str(e)}")
                except Exception as e:
                    paragraph.add_run(f"Error al insertar la imagen: {str(e)}")
            # Insertar salto de página
            doc.add_page_break()
    
    # Función para agregar la página del índice con marcador
    def agregar_indice_con_marcador(doc):
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        ReporteAnexo2.configurar_pagina(new_section)
        # Añadir el título "Índice" centrado y formateado
        titulo = doc.add_paragraph("ÍNDICE")
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = titulo.runs[0]
        run.bold = True
        run.font.size = Pt(14)
        run.font.underline = False  # Sin subrayado
        run.font.color.rgb = RGBColor(0, 0, 0)  # Color negro
        
        # Insertar marcador donde se añadirá la tabla de contenido
        toc_placeholder = doc.add_paragraph()
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), "0")
        bookmark_start.set(qn("w:name"), "TOCPlaceholder")
        toc_placeholder._p.append(bookmark_start)
        
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), "0")
        toc_placeholder._p.append(bookmark_end)
        doc.add_page_break()

    # Función para agregar el contenido de la quinta página con ajustes adicionales y alineación justificada
    def agregar_contenido_quinta_pagina(doc,respuesta):    
        if respuesta:
            ReporteAnexo2.agregar_titulo_1(doc,"INTRODUCCIÓN",WD_PARAGRAPH_ALIGNMENT.LEFT)
            # Objetivo
            ReporteAnexo2.agregar_titulo_1(doc,'OBJETIVO',WD_PARAGRAPH_ALIGNMENT.LEFT)
            ReporteAnexo2.justificar_parrafo(doc,respuesta[15])
            # Finalidad
            ReporteAnexo2.agregar_titulo_1(doc,'FINALIDAD',WD_PARAGRAPH_ALIGNMENT.LEFT)
            ReporteAnexo2.justificar_parrafo(doc,respuesta[16])

            # Ámbito de Aplicación
            ReporteAnexo2.agregar_titulo_1(doc,'ÁMBITO DE APLICACIÓN',WD_PARAGRAPH_ALIGNMENT.LEFT)
            ReporteAnexo2.justificar_parrafo(doc,respuesta[17])

            # Reporte Mensual con viñetas y justificado
            ReporteAnexo2.agregar_titulo_1(doc,'REPORTE MENSUAL',WD_PARAGRAPH_ALIGNMENT.LEFT)
            ReporteAnexo2.agregar_parrafos_con_vinetas(doc,respuesta[18])

            if respuesta[19] and respuesta[19].strip():
                anexo1 = doc.add_paragraph(f"ANEXO 2: {respuesta[19]}")
                anexo1.style = 'Normal'
                anexo1.paragraph_format.left_indent = Inches(0.5)
                anexo1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            # Insertar salto de página
            doc.add_page_break()
    
    ################### ANEXO 2##############################
    def generar_tabla_resumen_ejecutivo_A2(doc, componentes):
        if componentes:
            idcomponentes = [compo[0] for compo in componentes]
            respuesta = ReporteController.ctrlObtenerTablaResumenEjecutivoAnexo2(idcomponentes)
            if respuesta:
                # Inicializar el contador de letras (a., b., c., ...)
                letra_actual = 97  # Código ASCII de 'a'
                for fila in respuesta:
                    # Si hay más de una fila, añadir un título con viñeta alfabética
                    if len(respuesta) > 1:
                        letra = chr(letra_actual)
                        ReporteAnexo2.agregar_viñeta_letras_negrita(doc, fila[0], letra)
                        letra_actual += 1  # Incrementar letra para la siguiente tabla
                    # Crear una nueva tabla con 3 filas y 12 columnas
                    table = doc.add_table(rows=3, cols=12)
                    table.style = 'Table Grid'
                    table.allow_autofit = True
                    # Fila 1: DEPÓSITO - 6 columnas combinadas cada celda
                    cell_00 = table.cell(0, 0)
                    cell_00.text = fila[4]  # DEPÓSITO:
                    ReporteAnexo2.set_cell_background_color(table.cell(0, 0), "F0F0F0")
                    cell_00.merge(table.cell(0, 5))  # Combinar primeras 6 columnas
                    cell_01 = table.cell(0, 6)
                    cell_01.text = fila[5]  # DEPÓSITO DE RELAVE
                    cell_01.merge(table.cell(0, 11))  # Combinar últimas 6 columnas
                    # Fila 2: PERIODO
                    cell_10 = table.cell(1, 0)
                    cell_10.text = fila[6]  # PERIODO
                    ReporteAnexo2.set_cell_background_color(table.cell(1, 0), "F0F0F0")
                    cell_10.merge(table.cell(1, 5))  # Combinar primeras 6 columnas
                    cell_11 = table.cell(1, 6)
                    cell_11.text = fila[7]  # SETIEMBRE 2024
                    cell_11.merge(table.cell(1, 11))  # Combinar últimas 6 columnas
                    # Fila 3: Declaración (primera) - 8 columnas combinadas + 4 individuales
                    cell_20 = table.cell(2, 0)
                    cell_20.text = fila[8]  # Primera declaración
                    cell_20.merge(table.cell(2, 7))  # Combinar las primeras 8 columnas
                    table.cell(2, 8).text = "SI"
                    table.cell(2, 9).text = "X" if fila[9] == "SI" else "(\u00A0)"
                    table.cell(2, 10).text = "NO"
                    table.cell(2, 11).text = "X" if fila[9] == "NO" else "(\u00A0)"
                    # Aplicar color a las celdas con "X" en SI o NO
                    if fila[9] == "SI":
                        ReporteAnexo2.set_cell_background_color(table.cell(2, 8), "C6E0B4")
                        ReporteAnexo2.set_cell_background_color(table.cell(2, 9), "C6E0B4")
                    elif fila[9] == "NO":
                        ReporteAnexo2.set_cell_background_color(table.cell(2, 10), "C6E0B4")
                        ReporteAnexo2.set_cell_background_color(table.cell(2, 11), "C6E0B4")
                    # Aplicar tamaño de texto global (10 puntos) a todas las celdas
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.runs[0].font.size = Pt(10)
                    # Aplicar estilos a la fila 3
                    for cell in table.row_cells(2)[8:12]:  # Columnas SI, X, NO, X
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cell_20.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for paragraph in cell_20.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        paragraph.runs[0].font.size = Pt(10)
                    # Centramos la tabla y agregamos espacio al final
                    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Separar las tablas con un espacio
                    doc.add_paragraph("")
    
    def generar_tabla_instrumentacion(doc, componentes):
        if componentes:
            idcomponentes = [compo[0] for compo in componentes]
            respuesta = ReporteController.ctrlObtenerInstrumentacionAnexo2(idcomponentes)
            if respuesta:
                # Agrupar las filas según el valor de la tercera columna (índice 2)
                grupos = {}
                for fila in respuesta:
                    grupo = fila[2]  # Usamos el valor idcomponente
                    if grupo not in grupos:
                        grupos[grupo] = []
                    grupos[grupo].append(fila)
                # Inicializar la letra para viñetas en minúscula si hay más de un grupo
                agregar_letra = len(grupos) > 1
                letra = "a"
                # Iterar sobre los grupos y generar tablas
                for grupo, filas in grupos.items():
                    # Agregar título si hay más de un grupo
                    if agregar_letra:
                        titulo = filas[0][0]
                        ReporteAnexo2.agregar_viñeta_letras_negrita(doc, titulo, letra)
                        letra = chr(ord(letra) + 1)
                    # Crear la tabla con 12 columnas
                    table = doc.add_table(rows=1, cols=12)
                    table.style = 'Table Grid'
                    # Configuración de los encabezados
                    encabezados = [
                        ("Descripción", 1),
                        ("Autorizado", 4),
                        ("Adicional", 4),
                        ("Total", 2),
                        ("Frecuencia de Monitoreo", 1)
                    ]
                    # Pintar el encabezado completo
                    for cell in table.rows[0].cells:
                        ReporteAnexo2.set_cell_background_color(cell, "F0F0F0")
                    # Crear fila de encabezados combinados
                    fila_encabezado = table.rows[0]
                    col_actual = 0
                    for texto, span in encabezados:
                        cell = fila_encabezado.cells[col_actual]
                        cell.text = texto
                        if span > 1:
                            cell.merge(fila_encabezado.cells[col_actual + span - 1])
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            paragraph.runs[0].font.size = Pt(10)
                            paragraph.runs[0].font.bold = True
                        col_actual += span
                    # Insertar las filas de datos
                    for fila in filas:
                        # Agregar una nueva fila
                        row = table.add_row().cells
                        # Asignar los valores
                        row[0].text = fila[3]
                        row[1].text = 'Cantidad'
                        row[2].text = fila[4]
                        row[3].text = 'Operativo'
                        row[4].text = fila[5]
                        row[5].text = 'Cantidad'
                        row[6].text = fila[6]
                        row[7].text = 'Operativo'
                        row[8].text = fila[7]
                        row[9].text = 'Cantidad'
                        row[10].text = fila[8]
                        row[11].text = fila[9]
                        # Aplicar color de fondo a las dos celdas del apartado "Total"
                        ReporteAnexo2.set_cell_background_color(row[9], "C6E0B4")  # Celda "Cantidad" en Total
                        ReporteAnexo2.set_cell_background_color(row[10], "C6E0B4")  # Celda con el valor en Total
                        # Estilo y alineación para cada celda
                        for i, cell in enumerate(row):
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if i > 0 else WD_PARAGRAPH_ALIGNMENT.LEFT
                                paragraph.runs[0].font.size = Pt(9)
                    # Centramos la tabla y agregamos espacio al final
                    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Separar las tablas con un espacio
                    doc.add_paragraph("")
    
    def interpretacionValoresRegistrados(doc, componentes, idproyecto, fecha_inicio_formateada, fecha_fin_formateada):
        if componentes:
            if len(componentes) > 1:
                contador_grupo = 1
                for compon in componentes:
                    idcomponente = compon[0]
                    # Título del grupo principal (ej. 3.1, 3.2, 3.3...) y se agrega la ubicación (ubicación de la primera fila)
                    grupo_titulo = f"3.{contador_grupo} {compon[2]}"
                    # Agregar el título del grupo principal al documento en negrita
                    p = doc.add_paragraph()
                    p.add_run(grupo_titulo).bold = True
                    # Subgrupo numeración (ej. 3.1, 3.2, ...)
                    contador_subgrupo = 1
                    for item in [1, 2, 3, 4, 8, 5, 6, 7]: # "P Casagrande", "Cuerda Vibrante", "Hitos Topográficos", "Inclinómetros", "TDR", "Celdas", "Satelital", "Acelerógrafos"
                        ubicaciones = ReporteController.ctrlObtenerInterpretacionValoresA2(item, idcomponente)
                        if ubicaciones:
                            for fila in ubicaciones:
                                # Numeración de subgrupo (ej. 3.1.1, 3.1.2, ...)
                                subgrupo_titulo = f"3.{contador_grupo}.{contador_subgrupo} {fila[3]}"  # Columna 4: Nombre del subgrupo
                                p = doc.add_paragraph()
                                p.add_run(subgrupo_titulo).bold = True
                                # Agregar los tres puntos dentro del subgrupo
                                p = doc.add_paragraph()
                                p.add_run(f"a. Ubicación").bold = True
                                # Insertar la imagen BLOB en el centro de la página
                                paragraph = doc.add_paragraph()
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                run = paragraph.add_run()                                
                                try:
                                    # Convertir el blob a un archivo en memoria y agregarlo como imagen
                                    image_stream = BytesIO(fila[4])
                                    run.add_picture(image_stream, width=Inches(5)) 
                                except ZeroDivisionError:
                                    try:
                                        img_stream = ReporteAnexo2.reparar_imagen(image_stream)
                                        # Ajustar el tamaño de la imagen al ancho deseado
                                        run.add_picture(img_stream, width=Inches(5)) 
                                    except Exception as e:
                                        paragraph.add_run(f"Error al insertar la imagen: {str(e)}")
                                except Exception as e:
                                    paragraph.add_run(f"Error al insertar la imagen: {str(e)}")
                    
                                # Umbrales
                                p = doc.add_paragraph()
                                p.add_run(f"b. Umbrales de Alerta").bold = True
                                if fila[5] == 1: # casagrande
                                    ReporteAnexo2.generarTablaUmbrales(doc, idcomponente, idproyecto, 'umbral_piezometro', "PIEZOMETROMANUAL")
                                    ReporteAnexo2.generar_tabla_resumen_piezometros(doc,idcomponente, "Casagrande", idproyecto, fecha_inicio_formateada, fecha_fin_formateada)
                                    tipo_equipo = "Piezometromanual"
                                elif fila[5] == 2: # cuerda
                                    ReporteAnexo2.generarTablaUmbrales(doc,idcomponente, idproyecto, 'umbral_piezometro', "PIEZOMETROCUERDA")
                                    ReporteAnexo2.generar_tabla_resumen_piezometros(doc,idcomponente, "Cuerda", idproyecto, fecha_inicio_formateada, fecha_fin_formateada)
                                    tipo_equipo = "Piezometrocuerda"
                                elif fila[5] == 3: # prismas
                                    ReporteAnexo2.generarTablaUmbrales(doc,idcomponente,idproyecto,'umbral_prisma')
                                    ReporteAnexo2.generar_tabla_resumen_prismas(doc, idproyecto, idcomponente,fecha_inicio_formateada,fecha_fin_formateada)
                                    tipo_equipo = "Prisma"
                                elif fila[5] == 4: # inclinómetros
                                    ReporteAnexo2.generarTablaUmbrales(doc,idcomponente,idproyecto,'umbral_inclinometro')
                                    ReporteAnexo2.generar_tabla_resumen_inclinometros(doc, idproyecto, idcomponente, fecha_inicio_formateada, fecha_fin_formateada)
                                    tipo_equipo = "Inclinometro"
                                elif fila[5] == 5: # celdas
                                    ReporteAnexo2.generarTablaUmbrales(doc,idcomponente,idproyecto,'umbral_celda')
                                    ReporteAnexo2.generar_tabla_resumen_celdas(doc, idproyecto,idcomponente, fecha_inicio_formateada, fecha_fin_formateada)
                                    tipo_equipo = "Celda"
                                elif fila[5] == 6: # satelitales
                                    tipo_equipo = "Satelital"
                                    # ReporteAnexo2.generar_tabla_umbrales_satelitales(doc, fila[3], tipo_equipo, "anexo2")
                                elif fila[5] == 7: # acelerografos
                                    ReporteAnexo2.generar_tabla_umbrales_acelerografos(doc, idproyecto, idcomponente)
                                    tipo_equipo = "Acelerografo"
                                elif fila[5] == 8: # TDR
                                    # umbralestdr = UmbralController.ctrlObtenerUmbralSondajestdr(idproyecto)
                                    # ReporteAnexo2.generar_tabla_umbrales_sondajestdr(doc, umbralestdr)
                                    # ReporteAnexo2.generar_tabla_resumen_sondajestdr(doc, idproyecto, idcomponente, tiposubcomponente, umbralestdr, fecha_inicio_formateada, fecha_fin_formateada)
                                    tipo_equipo = "Sondajestdr"
                                # graficas reporte
                                p = doc.add_paragraph()
                                p.add_run(f"c. Evaluación de monitoreo").bold = True
                                imagenesreporte = ReporteController.ctrlObtenerImagenesGraficasReporte(idcomponente, tipo_equipo, "ANEXO2")
                                if imagenesreporte:
                                    for images in imagenesreporte:
                                        p = doc.add_paragraph()
                                        p.add_run(images[5])
                                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                        # Insertar la imagen BLOB en el centro de la página
                                        paragraph = doc.add_paragraph()
                                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                        run = paragraph.add_run()
                                        try:
                                        # Convertir el BLOB de imagen (bytes) en un flujo de bytes
                                            image_stream = BytesIO(images[4])
                                            # Insertar la imagen desde el flujo de bytes
                                            run.add_picture(image_stream, width=Inches(5))  # Ajustar el tamaño según sea necesario
                                        except ZeroDivisionError:
                                            try:
                                                img_stream = ReporteAnexo2.reparar_imagen(image_stream)
                                                # Insertar la imagen desde el flujo de bytes
                                                run.add_picture(img_stream, width=Inches(5))  # Ajustar el tamaño según sea necesario
                                            except Exception as e:
                                                paragraph.add_run(f"Error al insertar imagenes: {str(e)}")
                                        except Exception as e:
                                            paragraph.add_run(f"Error al insertar imagenes: {str(e)}")
                                        p = doc.add_paragraph()
                                        p.add_run(images[6])
                                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                contador_subgrupo += 1
                    # Incrementamos el contador del grupo principal
                    contador_grupo += 1
            else:
                for compon in componentes:
                    idcomponente = compon[0]
                    # tiposubcomponente = compon[4]
                    contador_subgrupo = 1
                    for item in [1, 2, 3, 4, 8, 5, 6, 7]:
                        ubicaciones = ReporteController.ctrlObtenerInterpretacionValoresA2(item, idcomponente)
                        if ubicaciones:
                            for fila in ubicaciones:
                                # Título del subgrupo principal (ej. 3.1, 3.2, 3.3...)
                                subgrupo_titulo = f"3.{contador_subgrupo} {fila[3]}"  # Usamos el valor de la columna 4 para el nombre del subgrupo
                                p = doc.add_paragraph()
                                p.add_run(subgrupo_titulo).bold = True
                                # Agregar los tres puntos dentro del subgrupo
                                p = doc.add_paragraph()
                                p.add_run(f"a. Ubicación").bold = True
                                # Insertar la imagen BLOB en el centro de la página
                                paragraph = doc.add_paragraph()
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                run = paragraph.add_run()                                
                                try:
                                    # Convertir el BLOB de imagen (bytes) en un flujo de bytes
                                    image_stream = BytesIO(fila[4])
                                    # Insertar la imagen desde el flujo de bytes
                                    run.add_picture(image_stream, width=Inches(5))  # Ajustar el tamaño según sea necesario
                                except ZeroDivisionError:
                                    try:
                                        img_stream = ReporteAnexo2.reparar_imagen(image_stream)
                                            # Insertar la imagen desde el flujo de bytes
                                        run.add_picture(img_stream, width=Inches(5))  # Ajustar el tamaño según sea necesario
                                    except Exception as e:
                                        paragraph.add_run(f"Error al insertar imagenes: {str(e)}")
                                except Exception as e:
                                    paragraph.add_run(f"Error al insertar imagenes: {str(e)}")
                                    
                                # umbrales
                                p = doc.add_paragraph()
                                p.add_run(f"b. Umbrales de Alerta").bold = True
                                p = doc.add_paragraph()
                                if fila[5] == 1: # casagrande
                                    ReporteAnexo2.generarTablaUmbrales(doc, idcomponente, idproyecto, 'umbral_piezometro', "PIEZOMETROMANUAL")
                                    ReporteAnexo2.generar_tabla_resumen_piezometros(doc,idcomponente, "Casagrande", idproyecto, fecha_inicio_formateada, fecha_fin_formateada)
                                    tipo_equipo = "Piezometrocasagrande"
                                elif fila[5] == 2: # cuerda
                                    ReporteAnexo2.generarTablaUmbrales(doc, idcomponente, idproyecto, 'umbral_piezometro', "PIEZOMETROCUERDA")
                                    ReporteAnexo2.generar_tabla_resumen_piezometros(doc,idcomponente, "Cuerda", idproyecto, fecha_inicio_formateada, fecha_fin_formateada)
                                    tipo_equipo = "Piezometrocuerda"
                                elif fila[5] == 3: # prismas
                                    ReporteAnexo2.generarTablaUmbrales(doc,idcomponente,idproyecto,'umbral_prisma')
                                    ReporteAnexo2.generar_tabla_resumen_prismas(doc, idproyecto, idcomponente, fecha_inicio_formateada,fecha_fin_formateada)
                                    tipo_equipo = "Prisma"
                                elif fila[5] == 4: # inclinómetros
                                    ReporteAnexo2.generarTablaUmbrales(doc,idcomponente,idproyecto,'umbral_inclinometro')
                                    ReporteAnexo2.generar_tabla_resumen_inclinometros(doc, idproyecto, idcomponente, fecha_inicio_formateada, fecha_fin_formateada)
                                    tipo_equipo = "Inclinometro"
                                elif fila[5] == 5: # celdas
                                    ReporteAnexo2.generarTablaUmbrales(doc,idcomponente,idproyecto,'umbral_celda')
                                    ReporteAnexo2.generar_tabla_resumen_celdas(doc, idproyecto, idcomponente,fecha_inicio_formateada, fecha_fin_formateada)
                                    tipo_equipo = "Celda"
                                elif fila[5] == 6: # satelitales
                                    tipo_equipo = "Satelital"
                                    # ReporteAnexo2.generar_tabla_umbrales_satelitales(doc, fila[3], tipo_equipo, "anexo2")
                                elif fila[5] == 7: # acelerografos
                                    ReporteAnexo2.generar_tabla_umbrales_acelerografos(doc, idproyecto,idcomponente)
                                    tipo_equipo = "Acelerografo"
                                elif fila[5] == 8: # TDR
                                    # umbralestdr = UmbralController.ctrlObtenerUmbralSondajestdr(idproyecto)
                                    # ReporteAnexo2.generar_tabla_umbrales_sondajestdr(doc, umbralestdr)
                                    # ReporteAnexo2.generar_tabla_resumen_sondajestdr(doc, idproyecto, idcomponente, tiposubcomponente, umbralestdr, fecha_inicio_formateada, fecha_fin_formateada)
                                    tipo_equipo = "Sondajetdr"
                                # mostrar imagenes de las gráficas
                                p = doc.add_paragraph()
                                p.add_run(f"c. Evaluación de monitoreo").bold = True
                                imagenesreporte = ReporteController.ctrlObtenerImagenesGraficasReporte(idcomponente, tipo_equipo, "ANEXO2")
                                if imagenesreporte:
                                    for images in imagenesreporte:
                                        p = doc.add_paragraph()
                                        p.add_run(images[5]) #titulo
                                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                        # Insertar la imagen BLOB en el centro de la página
                                        paragraph = doc.add_paragraph()
                                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                        run = paragraph.add_run()                                        
                                        try:
                                            # Convertir el BLOB de imagen (bytes) en un flujo de bytes
                                            image_stream = BytesIO(images[4]) #imagen
                                            # Insertar la imagen desde el flujo de bytes
                                            run.add_picture(image_stream, width=Inches(5))  # Ajustar el tamaño según sea necesario
                                        except ZeroDivisionError:
                                            try:
                                                img_stream = ReporteAnexo2.reparar_imagen(image_stream)
                                                    # Insertar la imagen desde el flujo de bytes
                                                run.add_picture(img_stream, width=Inches(5))  # Ajustar el tamaño según sea necesario
                                            except Exception as e:
                                                paragraph.add_run(f"Error al insertar imagenes: {str(e)}")
                                        except Exception as e:
                                            paragraph.add_run(f"Error al insertar imagenes: {str(e)}")
                                    
                                        p = doc.add_paragraph()
                                        p.add_run(images[6]) 
                                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                contador_subgrupo += 1
    
    def generar_tabla_observaciones_A2(doc, componentes):
        if componentes:
            idcomponentes = [compo[0] for compo in componentes]
            respuesta = ReporteController.ctrlObtenerObservacionesAnexo2(idcomponentes)
            if respuesta:
                # Inicializar la letra para las viñetas
                letra = "a"
                # Agrupar las filas según un criterio (e.g., columna 2 en este caso)
                grupos = {}
                for fila in respuesta:
                    grupo = fila[2]  # Usar la columna 2 como criterio de agrupación
                    if grupo not in grupos:
                        grupos[grupo] = []
                    grupos[grupo].append(fila)
                # Verificar si hay más de un grupo
                agregar_letra = len(grupos) > 1
                # Iterar sobre los grupos y generar tablas para cada grupo
                for grupo, filas in grupos.items():
                    # Si hay varios grupos, agregar una viñeta con la letra correspondiente
                    if agregar_letra:
                        titulo = filas[0][0]
                        ReporteAnexo2.agregar_viñeta_letras_negrita(doc, titulo, letra)
                        letra = chr(ord(letra) + 1)  # Pasar a la siguiente letra
                    # Crear la tabla con 1 fila para la cabecera y 6 columnas
                    table = doc.add_table(rows=1, cols=6)
                    table.style = 'Table Grid'
                    # Configuración de las celdas de encabezado
                    table.cell(0, 0).text = "Descripción"
                    table.cell(0, 1).text = "Condición Actual"
                    table.cell(0, 5).text = "Comentarios"
                    # Pintar el encabezado completo
                    for cell in table.rows[0].cells:
                        ReporteAnexo2.set_cell_background_color(cell, "F0F0F0")
                    # Combinar las celdas de "Condición Actual"
                    table.cell(0, 1).merge(table.cell(0, 4))
                    # Alinear y ajustar las celdas de encabezado
                    for cell in table.rows[0].cells:
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            if paragraph.runs:
                                paragraph.runs[0].font.size = Pt(10)
                    # Iterar sobre las filas de datos en el grupo
                    for row_data in filas:
                        tipo_fila = row_data[-1]  # Última columna que define el tipo (1 o 2)
                        row = table.add_row().cells  # Crear nueva fila
                        if tipo_fila == 1:
                            # Fila tipo 1: "SI/NO" con comentarios
                            descripcion = row_data[3]
                            condicion = row_data[4]
                            comentario = row_data[7]
                            row[0].text = descripcion
                            row[1].text = "SI"
                            row[2].text = "(X)" if condicion == "SI" else "()"
                            row[3].text = "NO"
                            row[4].text = "(X)" if condicion == "NO" else "()"
                            row[5].text = comentario
                            # Pintar celdas según la condición
                            if "(X)" in row[2].text:
                                ReporteAnexo2.set_cell_background_color(row[1], "C6E0B4")
                                ReporteAnexo2.set_cell_background_color(row[2], "C6E0B4")
                            if "(X)" in row[4].text:
                                ReporteAnexo2.set_cell_background_color(row[3], "C6E0B4")
                                ReporteAnexo2.set_cell_background_color(row[4], "C6E0B4")

                        elif tipo_fila == 2:
                            # Fila tipo 2: Descripción, Medidas de control, Plazo, Responsable
                            descripcion = row_data[3]
                            medidas_control = row_data[5]
                            plazo = row_data[6]
                            responsable = row_data[8]
                            row[0].text = descripcion
                            row[1].text = medidas_control
                            row[3].text = plazo
                            row[5].text = responsable
                            # Combinar celdas
                            row[1].merge(row[2])
                            row[3].merge(row[4])
                        # Alinear y ajustar tamaño de fuente en celdas
                        for j in range(6):
                            cell = row[j]
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            for paragraph in cell.paragraphs:
                                if j == 5:
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                else:
                                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if j in [1, 2, 3, 4] else WD_PARAGRAPH_ALIGNMENT.LEFT
                                if paragraph.runs:
                                    paragraph.runs[0].font.size = Pt(10)
                    # Centramos la tabla y agregamos espacio al final
                    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # Separar las tablas con un espacio
                    doc.add_paragraph("")
    
    def generarTablaUmbrales(doc, componente, proyecto, tabla, tipoequipo="GENERAL"):
        # Obtener los umbrales desde la fuente de datos
        umbrales = UmbralController.ctrObtenerUmbralesEquiposCP(proyecto, componente, tabla, tipoequipo)
        if umbrales:
            if tabla not in ('umbral_inclinometro', 'umbral_celda', 'umbral_piezometro'):
                orden = 9
            else:
                orden = 3
            # Organizar los umbrales por el último campo
            umbrales_organizados = {}
            for umbral in umbrales:
                clave = umbral[orden]
                if clave not in umbrales_organizados:
                    umbrales_organizados[clave] = []
                umbrales_organizados[clave].append(umbral)
            # Generar tablas en el documento Word
            for clave, umbrales_grupo in umbrales_organizados.items():
                # Obtener el título y la unidad
                if tabla == 'umbral_piezometro':
                    titulo, unidad = umbrales_grupo[0][0], 'msnm'
                elif tabla == 'umbral_inclinometro':
                    titulo, unidad = umbrales_grupo[0][0], 'm'
                else:
                    titulo, unidad = ReporteAnexo2.obtener_nombre_por_tipo(umbrales_grupo[0][9])
                    
                # Añadir título centrado
                titulo_paragraph = doc.add_paragraph(titulo)
                titulo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # Añadir tabla
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                # Añadir encabezados de la tabla y centrar el contenido
                hdr_cells = table.rows[0].cells
                headers = ['Condición', 'Riesgo', f'Rango ({unidad})', 'Acciones']
                for cell, header in zip(hdr_cells, headers):
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(10)  # Establecer tamaño de fuente a 10
                    # Aplicar color de fondo al encabezado
                    ReporteAnexo2.set_cell_background_color(cell, "F0F0F0")
                # Añadir filas de la tabla
                for i, umbral in enumerate(umbrales_grupo):
                    condicion, color, riesgo, valor, accion = umbral[4], umbral[5], umbral[6], umbral[7], umbral[8]
                    rango = f"<{valor}" if i == 0 else f"{umbrales_grupo[i-1][7]} - {valor}"
                    # Añadir fila con los datos
                    row_cells = table.add_row().cells
                    row_cells[0].text = condicion
                    row_cells[1].text = riesgo
                    row_cells[2].text = rango
                    row_cells[3].text = accion
                    # Centrar el contenido de las celdas y establecer tamaño de fuente a 10
                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                    # Aplicar color de fondo a la celda de condición
                    ReporteAnexo2.set_cell_background_color(row_cells[0], color)
                # Centrar la tabla y agregar espacio al final
                table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_paragraph("")
    
    def generar_tabla_umbrales_acelerografos(doc, idproyecto, idcomponente):
        # Crear la tabla con encabezado y filas de datos
        table = doc.add_table(rows=1, cols=5)  # Ajustar a 5 columnas
        table.style = 'Table Grid'

        # Encabezado de la tabla
        headers = ["Condición", "Riesgo", "Rango (Magnitud)", "Rango (Distancia)", "Acciones a Realizar"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

            # Pintar el encabezado con el color F0F0F0
            ReporteAnexo2.set_cell_background_color(cell, "F0F0F0")

        # Obtener los umbrales
        umbralesacelero = UmbralController.ctrlObtenerUmbralAcelerografos(idproyecto, idcomponente)
        if umbralesacelero:
            # Llenar la tabla con los datos de los umbrales
            for i, umbral in enumerate(umbralesacelero, start=1):
                row_cells = table.add_row().cells
                row_cells[0].text = umbral[3]  # Condición
                row_cells[1].text = umbral[4]  # Riesgo
                row_cells[2].text = f"M>{umbral[7]}"  # Rango (Magnitud)
                row_cells[3].text = f"Distancia menor a {umbral[6]}Km"  # Rango (Distancia)
                row_cells[4].text = umbral[8]  # Acciones a Realizar
                # Pintar la celda de condición con el color correspondiente
                ReporteAnexo2.set_cell_background_color(row_cells[0], umbral[5])
                # Ajustar el tamaño de la fuente de toda la fila
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
        # Centrar la tabla y agregar espacio al final
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("")
    
    def generar_tabla_umbrales_satelitales(doc, idcompo, instr, anexo):
        texto = """Se han establecido rangos de colores que están asociadas al desplazamiento, los valores de estos están diferenciados para el caso de análisis mensual. En la siguiente tabla, se presenta los niveles en función al deslazamiento. Los desplazamientos menores a 1 cm corresponden a una condición normal."""
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.add_run(texto)
        # Insertar la imagen BLOB en el centro de la página
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run()
        # Traer leyenda satelelital
        tipo = 1
        imagenesreporte = ReporteController.ctrlObtenerImagenesGraficasReporte(idcompo, instr, anexo, tipo)
        if imagenesreporte:                
            try:
                # Convertir el BLOB de imagen (bytes) en un flujo de bytes
                image_stream = BytesIO(imagenesreporte[0][3])
                # Insertar la imagen desde el flujo de bytes
                run.add_picture(image_stream, width=Inches(5))  # Ajustar el tamaño según sea necesario
            except ZeroDivisionError:
                try:
                    img_stream = ReporteAnexo2.reparar_imagen(image_stream)
                        # Insertar la imagen desde el flujo de bytes
                    run.add_picture(img_stream, width=Inches(5))  # Ajustar el tamaño según sea necesario
                except Exception as e:
                    paragraph.add_run(f"Error al insertar imagenes: {str(e)}")
            except Exception as e:
                paragraph.add_run(f"Error al insertar imagenes: {str(e)}")
        else:
            paragraph.add_run(" ")
    
    def generar_tabla_umbrales_sondajestdr(doc, umbralestdr):
        # Crear la tabla con encabezado y filas de datos
        table = doc.add_table(rows=4, cols=4)
        table.style = 'Table Grid'
        # Resto del encabezado
        table.cell(0, 0).text = "Condición"
        table.cell(0, 1).text = "Riesgo"
        table.cell(0, 2).text = "Rango (Impedancia)"
        table.cell(0, 3).text = "Acción por realizar"
        # Definir el tamaño y alineación para el encabezado
        for i in range(0, 4):
            table.cell(0, i).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in table.cell(0, i).paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
        # Filas de datos
        bajo, medio, alto = 15, 30, 30
        colores = ["C6E0B4", "FFFF00", "FF0000"]
        if umbralestdr:
            bajo, medio, alto = umbralestdr[5], umbralestdr[6], umbralestdr[6]
            colores = [umbralestdr[2], umbralestdr[3], umbralestdr[4]]
        data = [
            ["Normal", "Bajo", f"<{bajo}", "El monitoreo se realiza normalmente según el manual de mantenimiento."],
            ["Alerta", "Medio", f"{bajo} a {medio}", "Inspección y monitoreo más frecuente."],
            ["Peligro", "Alto", f">{medio}", "Retirar al personal y equipos del área. Actuar según manual de operación y el plan de contingencia de la instalación."]
        ]
        # Llenar las filas de datos
        for i, (condicion, riesgo, rango, accion) in enumerate(data):
            row = table.rows[i + 1]
            row.cells[0].text = condicion
            row.cells[1].text = riesgo
            row.cells[2].text = rango
            row.cells[3].text = accion
            # Aplicar color de fondo a la columna "Condición"
            ReporteAnexo2.set_cell_background_color(row.cells[1], colores[i])
            # Alinear el texto en cada celda
            for j, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                # Centramos horizontalmente las celdas de "Condición", "Riesgo" y "Rango del nivel freático"
                if j in [0, 1, 2]:  # Columnas de "Condición", "Riesgo" y "Rango del nivel freático"
                    alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                else:
                    alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # "Acción por realizar" alineada a la izquierda
                for paragraph in cell.paragraphs:
                    paragraph.alignment = alignment
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        # Centramos la tabla y agregamos espacio al final
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Separar las tablas con un espacio
        doc.add_paragraph("")  
    # REÚMENES
    def generar_tabla_resumen_piezometros(doc, idcomponente, tipo, idproyecto, fechaini, fechafin):
        # Crear la tabla con encabezado (2 filas) y sin filas de datos adicionales
        table = doc.add_table(rows=2, cols=7)  # No se añade una columna adicional para el ID
        table.style = 'Table Grid'

        # Encabezado de la tabla
        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 0).text = "Piezómetro"

        table.cell(0, 1).merge(table.cell(0, 2))
        table.cell(0, 1).text = "Fecha"
        table.cell(1, 1).text = "Inicial"
        table.cell(1, 2).text = "Final"

        # Fusionar las celdas de "Cota del terreno", "Cota del nivel del agua", "Estado" y "Condición según nivel de alerta"
        table.cell(0, 3).merge(table.cell(1, 3))
        table.cell(0, 3).text = "Cota del terreno\n(msnm)"
        table.cell(0, 4).merge(table.cell(1, 4))
        table.cell(0, 4).text = "Cota del nivel del agua\n(msnm)"
        table.cell(0, 5).merge(table.cell(1, 5))
        table.cell(0, 5).text = "Estado"
        table.cell(0, 6).merge(table.cell(1, 6))
        table.cell(0, 6).text = "Condición según nivel de alerta"

        # Color de fondo del encabezado
        for row in table.rows[:2]:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                ReporteAnexo2.set_cell_background_color(cell, "F0F0F0")

        # Obtener los datos según el tipo de piezómetro
        if tipo == "Cuerda":
            data = PiezometroController.ctrlObtenerResumenCuerdaReporte(idproyecto, idcomponente, fechaini, fechafin)
        else:
            data = PiezometroController.ctrlObtenerResumenCasagrandeReporte(idproyecto, idcomponente, fechaini, fechafin)

        if data:
            # Obtener los umbrales para los piezómetros
            ids = [item[0] for item in data]
            umbrales_piezometros = UmbralController.ctrObtenerUmbralesPiezometros(ids)

            # Verificar si umbrales_piezometros es None
            if umbrales_piezometros is None:
                umbrales_piezometros = []

            # Agregar filas de datos sin fila vacía extra
            for row_data in data:
                row = table.add_row().cells
                for i, text in enumerate(row_data[1:]):  # Comenzar desde el segundo elemento para omitir el ID
                    if i in [3, 4]:  # Cota del terreno y Cota del nivel del agua
                        row[i].text = f"{text:.3f}"
                    else:
                        row[i].text = str(text)
                    row[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for paragraph in row[i].paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Filtrar umbrales para el piezómetro actual
                piezometro_id = row_data[0]  # Asumiendo que el ID del piezómetro es el primer elemento
                umbrales = [umbral for umbral in umbrales_piezometros if umbral[2] == piezometro_id]
                # Comparar el valor en la posición 5 (Cota del nivel del agua) con los umbrales
                valor = abs(row_data[5])
                condicion = "Umbral Excedido"
                color = "#FF0000"  # Color rojo para "Umbral Excedido"
                if umbrales:
                    for i, umbral in enumerate(umbrales):
                        if valor < umbral[6]:
                            condicion = umbral[3]
                            color = umbral[4]
                            break
                        elif i == len(umbrales) - 1 and valor >= umbral[6]:
                            condicion = "Umbral Excedido"
                            color = "#FF0000"
                else:
                    condicion = "Sin umbral"
                    color = "F0F0F0"  # Color para "Sin umbral"

                row[6].text = condicion
                ReporteAnexo2.set_cell_background_color(row[6], color)

                # Asegurar que la celda de "Condición" esté centrada
                row[6].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in row[6].paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Aplicar el tamaño de fuente después de agregar todos los datos
            for row in table.rows[2:]:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
        else:
            # Agregar una fila que indique "Sin datos en las fechas seleccionadas"
            row = table.add_row().cells
            row[0].merge(row[-1])
            row[0].text = "Sin datos en las fechas seleccionadas"
            row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in row[0].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

        # Centramos la tabla y agregamos espacio al final
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Separar las tablas con un espacio
        doc.add_paragraph("")

    def generar_tabla_resumen_prismas(doc, idproyecto, idcomponente, fecha_inicial, fecha_final):
        # Crear la tabla con encabezado (2 filas) y sin filas de datos adicionales
        table = doc.add_table(rows=2, cols=6)
        table.style = 'Table Grid'

        # Encabezado de la tabla
        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 0).text = "Hito\nTopográfico"

        table.cell(0, 1).merge(table.cell(0, 2))
        table.cell(0, 1).text = "Fecha"
        table.cell(1, 1).text = "Inicial"
        table.cell(1, 2).text = "Final"

        table.cell(0, 3).merge(table.cell(1, 3))
        table.cell(0, 3).text = "Velocidad promedio\ncm/día"
        table.cell(0, 4).merge(table.cell(1, 4))
        table.cell(0, 4).text = "Estado"
        table.cell(0, 5).merge(table.cell(1, 5))
        table.cell(0, 5).text = "Condición según\nnivel de alerta"

        # Configurar alineación y color de fondo del encabezado
        for row in table.rows[:2]:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # Alineación vertical
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Alineación horizontal
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                # Aplicar color de fondo
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), "F0F0F0")
                cell._tc.get_or_add_tcPr().append(shading_elm)

        # Obtener los datos
        prismas = ReporteController.ctrlObtenerTablaResumenPrismas(idproyecto, idcomponente)
        if prismas:
            # Convertir los datos a un DataFrame
            df = pd.DataFrame(prismas, columns=['Nombre', 'Fecha', 'Valor1', 'Valor2', 'PenultimaColumna', 'Tipo'])
            # Convertir la columna de fechas a tipo datetime
            df['Fecha'] = pd.to_datetime(df['Fecha'])
            # Filtrar por fechas
            df_filtrado = df[(df['Fecha'] >= fecha_inicial) & (df['Fecha'] <= fecha_final)]
            if not df_filtrado.empty:
                # Agrupar por nombre y calcular el promedio de la penúltima columna
                resultado = df_filtrado.groupby('Nombre')['PenultimaColumna'].mean().reset_index()

                # Convertir las fechas a objetos datetime
                fecha_inicial_dt = datetime.strptime(fecha_inicial, '%Y-%m-%d')
                fecha_final_dt = datetime.strptime(fecha_final, '%Y-%m-%d')

                # Agregar las columnas de fecha inicial y fecha final
                resultado['FechaInicial'] = fecha_inicial_dt.strftime('%d/%m/%Y')
                resultado['FechaFinal'] = fecha_final_dt.strftime('%d/%m/%Y')

                # Renombrar la columna de promedio para mayor claridad
                resultado.rename(columns={'PenultimaColumna': 'PromedioPenultimaColumna'}, inplace=True)

                # Obtener los umbrales
                umbrales = UmbralController.ctrlObtenerUmbralesInstrumentacion(idproyecto, idcomponente, 'VI3D', 'umbral_prisma')

                # Llenar la tabla con los datos
                for _, row in resultado.iterrows():
                    cells = table.add_row().cells
                    cells[0].text = row['Nombre']
                    cells[1].text = row['FechaInicial']
                    cells[2].text = row['FechaFinal']
                    cells[3].text = f"{row['PromedioPenultimaColumna']:.3f}"
                    cells[4].text = "Operativo"

                    # Comparar el valor con los umbrales
                    valor = abs(row['PromedioPenultimaColumna'])
                    if not umbrales:
                        condicion = "sin umbral"
                        color = "#FFFFFF"  # Color blanco para "sin umbral"
                    else:
                        condicion = "Umbral Excedido"
                        color = "#FF0000"  # Color rojo para "Umbral Excedido"
                        for i, umbral in enumerate(umbrales):
                            if valor < umbral[6]:
                                condicion = umbral[3]
                                color = umbral[4]
                                break
                            elif i == len(umbrales) - 1 and valor >= umbral[6]:
                                condicion = "Umbral Excedido"
                                color = "#FF0000"

                    cells[5].text = condicion
                    # Aplicar color de fondo a la celda de condición
                    if condicion != "sin umbral":
                        ReporteAnexo2.set_cell_background_color(cells[5], color)

                    # Asegurar que todas las celdas estén centradas
                    for cell in cells:
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # Alineación vertical
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Alineación horizontal

            else:
                # Agregar una fila que indique "Sin datos en las fechas seleccionadas"
                row = table.add_row().cells
                row[0].merge(row[-1])
                row[0].text = "Sin datos en las fechas seleccionadas"
                row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # Alineación vertical
                for paragraph in row[0].paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Alineación horizontal
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        else:
            # Agregar una fila que indique "Sin datos en las fechas seleccionadas"
            row = table.add_row().cells
            row[0].merge(row[-1])
            row[0].text = "Sin datos en las fechas seleccionadas"
            row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # Alineación vertical
            for paragraph in row[0].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Alineación horizontal
                for run in paragraph.runs:
                    run.font.size = Pt(10)

        # Centramos la tabla y agregamos espacio al final
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Separar las tablas con un espacio
        doc.add_paragraph("")

    def generar_tabla_resumen_inclinometros(doc, idproyecto, idcomponente, fechainicial, fechafinal):
        # Crear la tabla con 2 filas iniciales (encabezado)
        table = doc.add_table(rows=2, cols=6)
        table.style = 'Table Grid'

        # Configurar el encabezado de la tabla
        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 0).text = "Inclinómetro"

        table.cell(0, 1).merge(table.cell(0, 2))
        table.cell(0, 1).text = "Fecha"
        table.cell(1, 1).text = "Inicial"
        table.cell(1, 2).text = "Final"

        table.cell(0, 3).merge(table.cell(1, 3))
        table.cell(0, 3).text = "Desplazamiento Acumulado\n(mm)"
        table.cell(0, 4).merge(table.cell(1, 4))
        table.cell(0, 4).text = "Estado"
        table.cell(0, 5).merge(table.cell(1, 5))
        table.cell(0, 5).text = "Condición según\nnivel de alerta"

        # Color de fondo del encabezado
        for row in table.rows[:2]:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                ReporteAnexo2.set_cell_background_color(cell, "F0F0F0")

        # Obtener los datos de la respuesta
        fechaini = fechainicial + " 00:00:00"
        fechafin = fechafinal + " 00:00:00"
        respuesta = InclinometroController.ctrlObtenerDAA_Inclinometro(idproyecto, idcomponente, 1000, fechaini, fechafin)

        # Verificar si hay datos en la respuesta
        if respuesta:
            ids = [item[0] for item in respuesta]
            umbrales_inclinometros = UmbralController.ctrObtenerUmbralesInclinometros(ids)
            for item in respuesta:
                # Desempaquetar los datos de cada fila
                inclinometro_id, nombre_inclinometro, fecha_inicial, fecha_final, desplazamiento = item

                # Agregar una nueva fila a la tabla
                row = table.add_row().cells

                # Formatear las fechas
                fecha_inicial_formatted = fecha_inicial.strftime('%d/%m/%Y')
                fecha_final_formatted = fecha_final.strftime('%d/%m/%Y')

                # Llenar las celdas con los datos
                row[0].text = str(nombre_inclinometro)  # Inclinómetro
                row[1].text = fecha_inicial_formatted  # Fecha inicial
                row[2].text = fecha_final_formatted  # Fecha final

                # Verificar si hay "Sin Lectura" en el desplazamiento
                if desplazamiento == "Sin Lectura":
                    row[3].text = "Sin Lectura"  # Desplazamiento Acumulado
                    row[4].text = "Falta fijar tubería"  # Estado
                    row[5].text = "N/A"  # Condición según nivel de alerta
                    ReporteAnexo2.set_cell_background_color(row[5], "F0F0F0")  # Color para "N/A"
                else:
                    row[3].text = f"{desplazamiento:.3f}"  # Desplazamiento Acumulado
                    row[4].text = "Operativo"  # Estado
                    # Comparar el valor del desplazamiento con los umbrales
                    valor = abs(desplazamiento)
                    condicion = "Umbral Excedido"
                    color = "#FF0000"  # Color rojo para "Umbral Excedido"
                    # Filtrar umbrales para el inclinómetro actual
                    if umbrales_inclinometros:
                        umbrales = [umbral for umbral in umbrales_inclinometros if umbral[2] == inclinometro_id]
                        for i, umbral in enumerate(umbrales):
                            if valor < umbral[6]:
                                condicion = umbral[3]
                                color = umbral[4]
                                break
                            elif i == len(umbrales) - 1 and valor >= umbral[6]:
                                condicion = "Umbral Excedido"
                                color = "#FF0000"
                    else:
                        condicion = "Sin umbral"
                        color = "F0F0F0"  # Color para "Sin umbral"

                    row[5].text = condicion  # Condición según nivel de alerta
                    ReporteAnexo2.set_cell_background_color(row[5], color)  # Pintar celda según umbral
                # Centrar el texto en las celdas de la fila
                for cell in row:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
        else:
            # Si no hay datos, agregar una fila indicando que no hay datos
            row = table.add_row().cells
            row[0].merge(row[-1])
            row[0].text = "Sin datos en las fechas seleccionadas"
            row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in row[0].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

        # Centrar la tabla y agregar espacio al final
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("")

    def generar_tabla_resumen_celdas(doc, idproyecto, idcomponente, fechainic, fechafina):
        # Crear la tabla con encabezado (2 filas) y sin filas de datos adicionales
        table = doc.add_table(rows=2, cols=6)
        table.style = 'Table Grid'

        # Encabezado de la tabla
        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 0).text = "Celda"

        table.cell(0, 1).merge(table.cell(0, 2))
        table.cell(0, 1).text = "Fecha"
        table.cell(1, 1).text = "Inicial"
        table.cell(1, 2).text = "Final"

        # Fusionar las celdas
        table.cell(0, 3).merge(table.cell(1, 3))
        table.cell(0, 3).text = "Asentamiento\n(m)"
        table.cell(0, 4).merge(table.cell(1, 4))
        table.cell(0, 4).text = "Estado"
        table.cell(0, 5).merge(table.cell(1, 5))
        table.cell(0, 5).text = "Condición según nivel de alerta"

        # Color de fondo del encabezado
        for row in table.rows[:2]:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                ReporteAnexo2.set_cell_background_color(cell, "F0F0F0")

        config = SoftwareConfiguracion.obtenerDataSoftware()
        confechas, positivo_negativo = config[16], config[19]
        # Obtener los datos de la respuesta
        info_celdas = InstrumentacionController.ctrlObtenerInstrumentacionComponente(idcomponente, 'CELDA')
        # Extraer los IDs de la ubicación 0
        ids_instrumentacion = [item[0] for item in info_celdas]
        datos = CeldaController.ctrlCalcularVelocidadMesReporte(idproyecto, idcomponente, ids_instrumentacion, fechainic, fechafina, confechas)

        # Convertir las fechas a objetos datetime
        fechaini = f"{fechainic} 00:00:00"
        fechafin = f"{fechafina} 00:00:00"
        fechainic_dt = datetime.strptime(fechaini, '%Y-%m-%d %H:%M:%S')
        fechafina_dt = datetime.strptime(fechafin, '%Y-%m-%d %H:%M:%S')

        # Filtrar los datos por el rango de fechas
        filtrado = [
            dato for dato in datos
            if fechainic_dt <= datetime.strptime(dato[2], '%Y-%m-%d %H:%M:%S') <= fechafina_dt
        ]

        # Agrupar por nombre y calcular el promedio de la columna 5
        resultado = []
        for nombre in set(dato[1] for dato in filtrado):
            grupo = [dato for dato in filtrado if dato[1] == nombre]
            if grupo:  # Verificar si el grupo no está vacío
                if positivo_negativo == 0:
                    promedio = sum(dato[8] for dato in grupo) / len(grupo)
                else:
                    promedio = sum(dato[5] for dato in grupo) / len(grupo)
            else:
                promedio = 'Sin lecturas'
            resultado.append((nombre, fechainic, fechafina, promedio))

        if resultado:
            umbrales = UmbralController.ctrlObtenerUmbralesInstrumentacion(idproyecto, idcomponente, 'V', 'umbral_celda')
            # Agregar filas con los datos de los promedios
            for nombre, fechainic, fechafina, promedio in resultado:
                # Formatear las fechas
                fechainic_formatted = fechainic_dt.strftime('%d/%m/%Y')
                fechafina_formatted = fechafina_dt.strftime('%d/%m/%Y')

                row = table.add_row().cells
                row_data = [nombre, fechainic_formatted, fechafina_formatted, f"{promedio:.3f}", "Operativo", "Normal"]
                for i, text in enumerate(row_data):
                    row[i].text = str(text)
                    row[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for paragraph in row[i].paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Comparar el valor en la posición 3 (promedio) con los umbrales y agregar el resultado en la última columna
                valor = abs(promedio)
                condicion = "Umbral Excedido"
                color = "#FF0000"  # Color rojo para "Umbral Excedido"
                for i, umbral in enumerate(umbrales):
                    if valor < umbral[6]:
                        condicion = umbral[3]
                        color = umbral[4]
                        break
                    elif i == len(umbrales) - 1 and valor >= umbral[6]:
                        condicion = "Umbral Excedido"
                        color = "#FF0000"

                row[5].text = condicion
                ReporteAnexo2.set_cell_background_color(row[5], color)

            # Aplicar el tamaño de fuente después de agregar todos los datos
            for row in table.rows[2:]:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
        else:
            # Agregar una fila que indique "Sin datos en las fechas seleccionadas"
            row = table.add_row().cells
            row[0].merge(row[-1])
            row[0].text = "Sin datos en las fechas seleccionadas"
            row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in row[0].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

        # Centramos la tabla y agregamos espacio al final
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Separar las tablas con un espacio
        doc.add_paragraph("")

    def generar_tabla_resumen_sondajestdr(doc, idproyecto, idcomponente, tiposubcomponente, umbralestdr, fecha_inicio_formateada, fecha_fin_formateada):
        table = doc.add_table(rows=2, cols=6)
        table.style = 'Table Grid'

        # Encabezado de la tabla
        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 0).text = "TDR"

        table.cell(0, 1).merge(table.cell(0, 2))
        table.cell(0, 1).text = "Fecha"
        table.cell(1, 1).text = "Inicial"
        table.cell(1, 2).text = "Final"

        table.cell(0, 3).merge(table.cell(1, 3))
        table.cell(0, 3).text = "Impedancia"
        table.cell(0, 4).merge(table.cell(1, 4))
        table.cell(0, 4).text = "Estado"
        table.cell(0, 5).merge(table.cell(1, 5))
        table.cell(0, 5).text = "Condición según\nnivel de alerta"

        # Centramos el texto del encabezado
        for row in table.rows[:2]:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        if tiposubcomponente == 0:
            sondajes = TDRController.ctrlListarSondajesTDR(idproyecto)
            tiposondaje = 1
        else:
            sondajes = TDRController.ctrlListarSondajestdrSubcomponente(idcomponente)
            tiposondaje = 2

        fecha_inicio = datetime.strptime(fecha_inicio_formateada, '%Y-%m-%d').strftime('%d/%m/%Y')
        fecha_fin = datetime.strptime(fecha_fin_formateada, '%Y-%m-%d').strftime('%d/%m/%Y')

        normal = 15
        if umbralestdr:
            normal = umbralestdr[5]

        if sondajes:
            # Llenar la tabla con los datos procesados
            for sondatdr in sondajes:
                if tiposondaje == 1:
                    ubicacion = sondatdr[2]  # nombre tdr
                else:
                    ubicacion = sondatdr[3]
                desplazamiento = f"< {normal}"  # Estático
                estado = "Operativo"    # Estático
                condicion = "Normal"    # Estático

                row_data = [ubicacion, fecha_inicio, fecha_fin, desplazamiento, estado, condicion]
                row = table.add_row().cells
                for i, text in enumerate(row_data):
                    row[i].text = str(text)
                    row[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for paragraph in row[i].paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                # Aplicar colores según el nivel de alerta
                color1, color2, color3 = "C6E0B4", "FFFF00", "FF0000"
                if umbralestdr:
                    color1, color2, color3 = umbralestdr[2], umbralestdr[3], umbralestdr[4]
                if row_data[5] == "Normal":
                    ReporteAnexo2.set_cell_background_color(row[5], color1)  # Verde
                elif row_data[5] == "Alerta":
                    ReporteAnexo2.set_cell_background_color(row[5], color2)  # Amarillo
                elif row_data[5] == "Peligro":
                    ReporteAnexo2.set_cell_background_color(row[5], color3)  # Rojo
        else:
            # Agregar una fila que indique "Sin datos en las fechas seleccionadas"
            row = table.add_row().cells
            row[0].merge(row[-1])
            row[0].text = "Sin datos en las fechas seleccionadas"
            row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in row[0].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

        # Centramos la tabla y agregamos espacio al final
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Separar las tablas con un espacio
        doc.add_paragraph("")

    def generar_tablaFirma(doc, responsable, fecha_completa):
        # Crear la tabla con 5 filas y 1 columna
        table = doc.add_table(rows=5, cols=1)
        table.style = 'Table Grid'  # Estilo de tabla

        # Ajustar el ancho de la columna para que la imagen se ajuste correctamente
        table.columns[0].width = Inches(2.5)  # Ajustar el ancho de la columna

        # Insertar la imagen de la firma en la primera celda
        cell_firma = table.cell(0, 0)
        cell_firma.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Añadir un párrafo con un carácter invisible antes de la imagen para espaciado
        paragraph_before = cell_firma.add_paragraph()
        paragraph_before.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_before.space_before = Pt(0)  # Ajustar el espaciado antes del párrafo a 0 puntos
        paragraph_before.space_after = Pt(0)   # Ajustar el espaciado después del párrafo a 0 puntos
        run_before = paragraph_before.add_run('\u00A0')  # Carácter invisible
        run_before.font.size = Pt(0.5)  # Fuente de 1 punto

        paragraph = cell_firma.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.space_before = Pt(0)  # Ajustar el espaciado antes del párrafo a 0 puntos
        paragraph.space_after = Pt(0)   # Ajustar el espaciado después del párrafo a 0 puntos
        run = paragraph.add_run()

        # Convertir el blob en una imagen y agregarla al documento
        if responsable:
            blob_firma = responsable[6]
            if blob_firma:
                try:
                    image_stream = BytesIO(blob_firma)  # Convertir el blob a un objeto BytesIO
                    # Ajustar el tamaño de la imagen al ancho deseado
                    run.add_picture(image_stream, width=Inches(1.5))
                except ZeroDivisionError:
                    try:
                        img_stream = ReporteAnexo2.reparar_imagen(image_stream)
                        # Ajustar el tamaño de la imagen al ancho deseado
                        run.add_picture(img_stream, width=Inches(1.5))
                    except Exception as e:
                        paragraph.add_run(f"Error al insertar la imagen: {str(e)}")
                except Exception as e:
                    paragraph.add_run(f"Error al insertar la imagen: {str(e)}")

            data = [
                responsable[2],
                f"DNI {responsable[4]}",
                f"CIP {responsable[5]}",
                f"FECHA: {fecha_completa}"
            ]
        else:
            data = [
                "Nombres y Apellidos",
                "DNI 04055799",
                "CIP 75616",
                f"FECHA: {fecha_completa}"
            ]

        # Añadir un párrafo con un carácter invisible después de la imagen para espaciado
        paragraph_after = cell_firma.add_paragraph()
        paragraph_after.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_after.line_spacing = Pt(0)  # Ajustar el espaciado entre líneas a 0 puntos
        run_after = paragraph_after.add_run('\u00A0')  # Carácter invisible
        run_after.font.size = Pt(2)  # Fuente de 1 punto

        # Llenar las celdas con el texto
        for i, text in enumerate(data):
            cell = table.cell(i + 1, 0)
            cell.text = text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.line_spacing = Pt(0)  # Ajustar el espaciado entre líneas a 0 puntos
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Ajuste de tamaño de fuente
                # Mantener el párrafo junto
                ReporteAnexo2.set_keep_together(paragraph)

        # Alinear la tabla al borde derecho de la página
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_page_break()
        # Agregar párrafos vacíos para centrar verticalmente el título
        for _ in range(10):  # Ajustar el número de párrafos según sea necesario
            doc.add_paragraph()

        # Crear un párrafo centrado con el título "Misceláneos PLANOS"
        titulo_paragraph = doc.add_paragraph()
        titulo_run = titulo_paragraph.add_run("Misceláneos\nPLANOS")
        titulo_run.bold = True  # Aplicar negrita
        titulo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Ajustar el tamaño de fuente del título
        titulo_run.font.size = Pt(14)  # Ajuste de tamaño de fuente
                    
    # Función para agregar el contenido de la quinta página con ajustes adicionales y alineación justificada
    def agregar_contenido_anexo2(idproyecto, componentes, doc, responsable, reporte, fecha_completa, fecha_inicio_formateada, fecha_fin_formateada):
        if componentes:
            ReporteAnexo2.agregar_titulo_1(doc, f"ANEXO 2: {reporte.upper()}", WD_PARAGRAPH_ALIGNMENT.LEFT)
            doc.add_paragraph()
            ReporteAnexo2.agregar_viñeta_numerica_negrita(doc,'1. RESÚMEN EJECUTIVO')
            ReporteAnexo2.generar_tabla_resumen_ejecutivo_A2(doc, componentes)
            ReporteAnexo2.agregar_viñeta_numerica_negrita(doc,'2. INSTRUMENTACIÓN GEOTÉCNICA')
            ReporteAnexo2.generar_tabla_instrumentacion(doc, componentes)
            ReporteAnexo2.agregar_viñeta_numerica_negrita(doc,'3. INTERPRETACIÓN DE VALORES REGISTRADOS')
            ReporteAnexo2.interpretacionValoresRegistrados(doc, componentes, idproyecto, fecha_inicio_formateada, fecha_fin_formateada)
            ReporteAnexo2.agregar_viñeta_numerica_negrita(doc,'4. OBSERVACIONES, MEDIDAS ADOPTADAS Y SEGUIMIENTO ')
            ReporteAnexo2.generar_tabla_observaciones_A2(doc, componentes)
            ReporteAnexo2.generar_tablaFirma(doc, responsable, fecha_completa)
    
    # Configuración principal del documento
    def configurar_documento_apa(proyectoid, fecha_completa, mes_anio, dia_mes_anio, fecha_inicio_formateada, fecha_fin_formateada):
        respuesta = ReporteController.ctrlListarDatosGeneralAnexos(proyectoid, 'Anexo2')
        ReporteAnexo2.doc = ReporteAnexo2.crear_documento()
        if respuesta:
            componentes = ReporteController.ctrlObtenerComponentes(proyectoid)
            ReporteAnexo2.configurar_fuente(ReporteAnexo2.doc)
            datos_generales = EmpresaController.ctrlObtenerDatosConfiguracionEmpresa()
            ReporteAnexo2.agregar_caratula(ReporteAnexo2.doc, datos_generales[5], respuesta, mes_anio)            
            ReporteAnexo2.agregar_encabezado(ReporteAnexo2.doc, datos_generales)
            ReporteAnexo2.agregar_pie_pagina(ReporteAnexo2.doc)
            responsable = EmpresaController.ctrlObtenerDatosConfiguracionResponsable(proyectoid)
            ReporteAnexo2.agregar_contenido_memorandum(ReporteAnexo2.doc, respuesta, responsable, datos_generales[1], dia_mes_anio)
            ReporteAnexo2.agregar_tercera_pagina(ReporteAnexo2.doc, respuesta, mes_anio)
            ReporteAnexo2.agregar_indice_con_marcador(ReporteAnexo2.doc)
            ReporteAnexo2.agregar_contenido_quinta_pagina(ReporteAnexo2.doc, respuesta)
            ReporteAnexo2.agregar_contenido_anexo2(proyectoid, componentes, ReporteAnexo2.doc, responsable, respuesta[19], fecha_completa, fecha_inicio_formateada, fecha_fin_formateada)

    def generarReporte(proyectoid, fechainicio, fechafinal):
        timezone = pytz.timezone("America/Lima")
        fecha_actual = datetime.now(timezone)
        mes_anio = MetodosGenerales.obtenerMesAnio(fechafinal)
        dia_mes_anio = MetodosGenerales.obtenerDiaMesAnio(fecha_actual)
        fecha_formateada = fecha_actual.strftime("%d/%m/%Y")
        fecha_inicio_formateada = fechainicio.toString("yyyy-MM-dd")
        fecha_fin_formateada = fechafinal.toString("yyyy-MM-dd")
        # actualizar indice
        ReporteAnexo2.configurar_documento_apa(proyectoid, fecha_formateada, mes_anio, dia_mes_anio, fecha_inicio_formateada, fecha_fin_formateada)
        doc_path = "modules/reportes/ANEXO2.docx"
        pdf_path = "modules/reportes/ANEXO2.pdf"
        ActualizarReporte.guardar_y_actualizar_indice_y_ajustar_tablas(doc_path,pdf_path,ReporteAnexo2.doc)
    
    def guardarInformacionGeneralAnexo2(main, idproyecto, imagencomponente, tipoanexo):
        if idproyecto:
            # PORTADA
            titulo_portada = main.findChild(QLineEdit, "input_titulo_portada_A2").text()
            subtitulo_portada = main.findChild(QLineEdit, "input_subtitulo_portada_A2").text()
            lugar_portada = main.findChild(QLineEdit, "input_lugar_portada_A2").text()
            autor_portada = main.findChild(QLineEdit, "input_autor_portada_A2").text()
            # DOCUMENTO
            tipo_documento = main.findChild(QLineEdit, "input_tipo_documento_A2").text()
            codigo_reporte = main.findChild(QLineEdit, "input_codigo_reporte_A2").text()
            destinatario = main.findChild(QPlainTextEdit, "input_destinatario_reporte_A2").toPlainText()
            remitente = main.findChild(QPlainTextEdit, "input_remitente_reporte_A2").toPlainText()
            asunto = main.findChild(QPlainTextEdit, "input_asunto_reporte_A2").toPlainText()
            descripcion = main.findChild(QPlainTextEdit, "input_descripcion_reporte_A2").toPlainText()
            # COMPONENTE
            tipo_reporte = main.findChild(QLineEdit, "input_tipo_reporte_A1").text()
            nombre_componente = main.findChild(QLineEdit, "input_componente_reporte_A2").text()
            # INTRODUCCIÓN
            objetivo_reporte = main.findChild(QPlainTextEdit, "input_objetivo_reporte_A2").toPlainText()
            finalidad_reporte = main.findChild(QPlainTextEdit, "input_finalidad_reporte_A2").toPlainText()
            ambito_reporte = main.findChild(QPlainTextEdit, "input_ambito_reporte_A2").toPlainText()
            detalles_reporte = main.findChild(QPlainTextEdit, "input_detalle_reporte_A2").toPlainText()
            titulo_anexo = main.findChild(QLineEdit, "input_titulo_anexo_A2").text()
            # Agrupar los datos en una lista
            datos = [titulo_portada, subtitulo_portada, lugar_portada, autor_portada, tipo_documento, codigo_reporte,
                     destinatario, remitente, asunto, descripcion, tipo_reporte, nombre_componente, imagencomponente,
                     objetivo_reporte, finalidad_reporte, ambito_reporte, detalles_reporte, titulo_anexo, tipoanexo]
            respuesta = ReporteController.ctrlGuardarDataGeneralAnexos(datos, idproyecto, tipoanexo)
            return respuesta
    
    def cargarDataFormulariosAnexoGeneral(main, idproyecto):
        general = ReporteController.ctrlListarDatosGeneralAnexos(idproyecto, "Anexo2")
        if general:
            main.findChild(QLineEdit, "input_titulo_portada_A2").setText(general[2])
            main.findChild(QLineEdit, "input_subtitulo_portada_A2").setText(general[3])
            main.findChild(QLineEdit, "input_lugar_portada_A2").setText(general[4])
            main.findChild(QLineEdit, "input_autor_portada_A2").setText(general[5])
            # DOCUMENTO
            main.findChild(QLineEdit, "input_tipo_documento_A2").setText(general[6])
            main.findChild(QLineEdit, "input_codigo_reporte_A2").setText(general[7])
            main.findChild(QPlainTextEdit, "input_destinatario_reporte_A2").setPlainText(general[8])
            main.findChild(QPlainTextEdit, "input_remitente_reporte_A2").setPlainText(general[9])
            main.findChild(QPlainTextEdit, "input_asunto_reporte_A2").setPlainText(general[10])
            main.findChild(QPlainTextEdit, "input_descripcion_reporte_A2").setPlainText(general[11])
            # COMPONENTE
            main.findChild(QLineEdit, "input_tipo_reporte_A2").setText(general[12])
            main.findChild(QLineEdit, "input_componente_reporte_A2").setText(general[13])
            if general[14]:
                pixmap = MetodosGenerales.convertir_blob_a_pixmap(general[14])
                scaled_pixmap = pixmap.scaled(120, 120, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                main.findChild(QLabel, "lb_imagen_componente_A2").setPixmap(scaled_pixmap)
            # INTRODUCCIÓN
            main.findChild(QPlainTextEdit, "input_objetivo_reporte_A2").setPlainText(general[15])
            main.findChild(QPlainTextEdit, "input_finalidad_reporte_A2").setPlainText(general[16])
            main.findChild(QPlainTextEdit, "input_ambito_reporte_A2").setPlainText(general[17])
            main.findChild(QPlainTextEdit, "input_detalle_reporte_A2").setPlainText(general[18])
            main.findChild(QLineEdit, "input_titulo_anexo_A2").setText(general[19])
    
    def guardarInformacionResumenEjecutivo(main, idproyecto):
        if idproyecto:
            fecha_hora_actual = datetime.now()
            fechahora = fecha_hora_actual.strftime("%Y-%m-%d %H:%M:%S") 
            idcomponente = main.findChild(QComboBox, "cb_componentes_anexos").currentData()
            descripcion_anexo = main.findChild(QPlainTextEdit, "input_descripcion_general_anexo_A2").toPlainText()
            componente_encabezado_a2 = main.findChild(QLineEdit, "input_componente_encabezado_anexo_A2").text()
            valor_componente_encabezado_a2 = main.findChild(QLineEdit, "input_valor_componente_encabezado_anexo_A2").text()
            periodo_encabezado_a2 = main.findChild(QLineEdit, "input_periodo_encabezado_anexo_A2").text()
            valor_periodo_encabezado_a2 = main.findChild(QLineEdit, "input_valor_periodo_encabezado_anexo_A2").text()
            interpreta_control_a2 = main.findChild(QPlainTextEdit, "input_interpretacion_monitoreo_anexo_A2").toPlainText()
            si_interpreta_a2 = main.findChild(QRadioButton, "rb_interpretacion_SI_anexo_A2")
            valor_interpreta_a2 = 'SI' if si_interpreta_a2.isChecked() else 'NO'
            # Enviar los datos al controlador
            valores = [descripcion_anexo, componente_encabezado_a2, valor_componente_encabezado_a2,
                periodo_encabezado_a2, valor_periodo_encabezado_a2, interpreta_control_a2, valor_interpreta_a2]
            respuesta = ReporteController.ctrlGuardarResumenEjecutivoAnexo2(valores, idcomponente)
            return respuesta
    
    def cargarInformacionResumenEjecutivoAnexo(main):
        idcomponente = main.findChild(QComboBox, "cb_componentes_anexos").currentData()
        resumen = ReporteController.ctrlObtenerResumenEjecutivoAnexo2(idcomponente)
        if resumen:
            main.findChild(QPlainTextEdit, "input_descripcion_general_anexo_A2").setPlainText(resumen[2])
            main.findChild(QLineEdit, "input_componente_encabezado_anexo_A2").setText(resumen[3])
            main.findChild(QLineEdit, "input_valor_componente_encabezado_anexo_A2").setText(resumen[4])
            main.findChild(QLineEdit, "input_periodo_encabezado_anexo_A2").setText(resumen[5])
            main.findChild(QLineEdit, "input_valor_periodo_encabezado_anexo_A2").setText(resumen[6])
            main.findChild(QPlainTextEdit, "input_interpretacion_monitoreo_anexo_A2").setPlainText(resumen[7])
            si_interpreta_a2 = main.findChild(QRadioButton, "rb_interpretacion_SI_anexo_A2")
            no_interpreta_a2 = main.findChild(QRadioButton, "rb_interpretacion_NO_anexo_A2")
            if resumen[8] == 'SI':
                si_interpreta_a2.setChecked(True)
            else:
                no_interpreta_a2.setChecked(True)
    
    def guardarDatosDinamicosAnexo2(widget_anexo2, componente):
        intrumentacion_a2 = ReporteAnexo2.obtener_valores_frame_instrumentacion_A2(widget_anexo2)
        if intrumentacion_a2:
            ReporteController.ctrlGuardarInstrumentacionGeotecnicaA2(componente, intrumentacion_a2)
        ubicacion_intrumentacion_a2 = ReporteAnexo2.obtener_valores_frame_ubicacion_intrumentacion_A2(widget_anexo2)
        if ubicacion_intrumentacion_a2:
            ReporteController.ctrlGuardarUbicacionesInstrumentacion(componente, ubicacion_intrumentacion_a2)
        observaciones_a2 = ReporteAnexo2.obtener_valores_frame_observaciones_A2(widget_anexo2)
        if observaciones_a2:
            ReporteController.ctrlGuardarObservacionesAnexo2(componente, observaciones_a2)
    
    def obtener_valores_frame_instrumentacion_A2(widget):
        # Buscar el layout principal del widget
        layout = widget.layout()
        # Acceder al frame3 (que está dentro del layout, en el índice 2)
        frame = layout.itemAt(0).widget()  # Cambiar índice según sea necesario
        # Verificar que frame3 exista
        if not frame:
            raise ValueError("No se encontró el frame3 en el índice especificado.")
        # Obtener el layout de frame3
        frame_layout = frame.layout()
        # Lista para almacenar los valores de las filas
        valores = []
        # Iterar a partir de la fila 3 (donde inician los datos)
        for i in range(3, frame_layout.rowCount()):
            row_values = []
            # Celda 0: QLineEdit (Descripción)
            description_input = frame_layout.itemAtPosition(i, 0).widget()
            description_text = description_input.text().strip() if description_input else ""
            # Si la descripción está vacía, no considerar la fila
            if not description_text:
                continue
            # Añadir descripción al registro
            row_values.append(description_text)
            # Celda 1: QLineEdit (Cantidad autorizada)
            authorized_quantity_input = frame_layout.itemAtPosition(i, 1).widget()
            row_values.append(authorized_quantity_input.text().strip() if authorized_quantity_input else "0")
            # Celda 2: QLineEdit (Operativo autorizado)
            authorized_operational_input = frame_layout.itemAtPosition(i, 2).widget()
            row_values.append(authorized_operational_input.text().strip() if authorized_operational_input else "0")
            # Celda 3: QLineEdit (Cantidad adicional)
            additional_quantity_input = frame_layout.itemAtPosition(i, 3).widget()
            row_values.append(additional_quantity_input.text().strip() if additional_quantity_input else "0")
            # Celda 4: QLineEdit (Operativo adicional)
            additional_operational_input = frame_layout.itemAtPosition(i, 4).widget()
            row_values.append(additional_operational_input.text().strip() if additional_operational_input else "0")
            # Celda 5: QComboBox (Frecuencia)
            frequency_combobox = frame_layout.itemAtPosition(i, 5).widget()
            row_values.append(frequency_combobox.currentText() if frequency_combobox else "")
            # Añadir los valores de esta fila a la lista
            valores.append(row_values)
        return valores
    
    def obtener_valores_frame_ubicacion_intrumentacion_A2(widget):
        # Buscar el layout principal del widget
        layout = widget.layout()
        # Asegurar que el layout existe
        if not layout:
            return []
        frame = layout.itemAt(1).widget()
        if not frame:
            return []
        # Obtener el layout dentro de frame4
        frame_layout = frame.layout()
        if not frame_layout:
            return []
        # Lista para almacenar los valores extraídos
        valores = []
        for row in range(2, frame_layout.rowCount()):
            instrumento_item = frame_layout.itemAtPosition(row, 0)
            imagen_item = frame_layout.itemAtPosition(row, 1)
            if instrumento_item and imagen_item:
                instrumento_line_edit = instrumento_item.widget()
                imagen_label = imagen_item.widget()
                if isinstance(instrumento_line_edit, QLineEdit) and isinstance(imagen_label, QLabel):
                    # Obtener el texto del instrumento
                    instrumento = instrumento_line_edit.text().strip()
                    row_id = getattr(instrumento_line_edit, "row_id", None)
                    tipo_instrumentacion = getattr(instrumento_line_edit, "tipo_instrumentacion", None)
                    if instrumento:  # Solo considerar filas con nombre de instrumento
                        pixmap = imagen_label.pixmap()
                        imagen_path = None
                        if pixmap:
                            imagen_path = getattr(imagen_label, "ruta_imagen", None)
                        # Solo agregar la fila si la imagen tiene un valor válido
                        if imagen_path and imagen_path != "Sin imagen":
                            valores.append({
                                "id": row_id,
                                "tipo_instrumentacion": tipo_instrumentacion,
                                "instrumento": instrumento,
                                "imagen": imagen_path
                            })
                        else:
                            valores.append({
                                "id": row_id,
                                "tipo_instrumentacion": tipo_instrumentacion,
                                "instrumento": instrumento,
                                "imagen": None
                            })
        return valores
    
    def obtener_valores_frame_observaciones_A2(widget):
        # Buscar el layout principal del widget
        layout = widget.layout()
        # Acceder al frame6 (que está dentro del layout, en el índice 5)
        frame = layout.itemAt(2).widget()  # Obtener el frame6 (índice 5)
        # Obtener el layout dentro de frame6
        frame_layout = frame.layout()
        # Lista para almacenar los valores de las filas
        valores = []
        # Recorremos las filas dinámicas, comenzando desde la fila 2 (porque las filas 0 y 1 son encabezados)
        for i in range(2, frame_layout.rowCount()):
            row_values = []
            # Celda 0: QPlainTextEdit (Descripción)
            description_widget = frame_layout.itemAtPosition(i, 0).widget()
            if not description_widget or not description_widget.toPlainText().strip():
                # Si la descripción está vacía, omitir esta fila
                continue
            descripcion = description_widget.toPlainText().strip()
            row_values.append(descripcion)  # Agregar Descripción
            # Determinar el tipo de fila (SI/NO o Comentarios Extendidos)
            condition_widget = frame_layout.itemAtPosition(i, 1).widget()
            if isinstance(condition_widget, QGroupBox):
                # Es una fila SI/NO
                radio_buttons = condition_widget.findChildren(QRadioButton)
                condicion_actual = next((radio.text() for radio in radio_buttons if radio.isChecked()), "")
                row_values.append(condicion_actual)  # Condición Actual
                row_values.append("")  # Medidas de Control vacío
                row_values.append("")  # Plazo vacío
                comments_widget = frame_layout.itemAtPosition(i, 3).widget()
                comentario = comments_widget.toPlainText().strip() if comments_widget else ""
                row_values.append(comentario)  # Comentarios
                row_values.append("")  # Responsable vacío
                row_values.append(1)  # Tipo de fila: 1 (SI/NO)
            else:
                # Es una fila de Comentarios Extendidos
                row_values.append("")  # Condición Actual vacío
                medidas_control_widget = frame_layout.itemAtPosition(i, 1).widget()
                medidas_control = medidas_control_widget.toPlainText().strip() if medidas_control_widget else ""
                row_values.append(medidas_control)  # Medidas de Control
                plazo_widget = frame_layout.itemAtPosition(i, 2).widget()
                plazo = plazo_widget.toPlainText().strip() if plazo_widget else ""
                row_values.append(plazo)  # Plazo
                comments_widget = frame_layout.itemAtPosition(i, 3).widget()
                responsable = comments_widget.toPlainText().strip() if comments_widget else ""
                row_values.append("")  # Comentarios vacío
                row_values.append(responsable)  # Responsable
                row_values.append(2)  # Tipo de fila: 2 (Comentarios Extendidos)
            # Añadir los valores de esta fila a la lista de valores
            valores.append(row_values)
        return valores
    
    def mostrar_mensaje(titulo, mensaje, icono):
        msg_box = QMessageBox()
        msg_box.setWindowTitle(titulo)
        msg_box.setText(mensaje)
        msg_box.setIcon(icono)
        msg_box.exec()

    def obtener_nombre_por_tipo(tipo):
        # Diccionario que mapea tipos a nombres y unidades
        tipo_a_info = {
            "3DA": {"nombre": "Desplazamiento Acum. 3D", "unidad": "m"},
            "3DI": {"nombre": "Desplazamiento Incr. 3D", "unidad": "m"},
            "2DA": {"nombre": "Desplazamiento Acum. 2D", "unidad": "m"},
            "2DI": {"nombre": "Desplazamiento Incr. 2D", "unidad": "m"},
            "SDA": {"nombre": "Desplazamiento Acum. SD", "unidad": "m"},
            "SDI": {"nombre": "Desplazamiento Incr. SD", "unidad": "m"},
            "DLA": {"nombre": "Desplazamiento Acum. L", "unidad": "m"},
            "DLI": {"nombre": "Desplazamiento Incr. L", "unidad": "m"},
            "DTA": {"nombre": "Desplazamiento Acum. T", "unidad": "m"},
            "DTI": {"nombre": "Desplazamiento Incr. T", "unidad": "m"},
            "DHA": {"nombre": "Desplazamiento Acum. H", "unidad": "m"},
            "DHI": {"nombre": "Desplazamiento Incr. H", "unidad": "m"},
            "DNA": {"nombre": "Desplazamiento Acum. N", "unidad": "m"},
            "DNI": {"nombre": "Desplazamiento Incr. N", "unidad": "m"},
            "DEA": {"nombre": "Desplazamiento Acum. E", "unidad": "m"},
            "DEI": {"nombre": "Desplazamiento Incr. E", "unidad": "m"},
            "DZA": {"nombre": "Desplazamiento Acum. Z", "unidad": "m"},
            "DZI": {"nombre": "Desplazamiento Incr. Z", "unidad": "m"},
            "VI3D": {"nombre": "Velocidad Incremental 3D", "unidad": "m"},
            "VA3D": {"nombre": "Velocidad Acumulada 3D", "unidad": "m"},
            "VI2D": {"nombre": "Velocidad Incremental 2D", "unidad": "m"},
            "VA2D": {"nombre": "Velocidad Acumulada 2D", "unidad": "m"},
            "VISD": {"nombre": "Velocidad Incremental SD", "unidad": "m"},
            "VASD": {"nombre": "Velocidad Acumulada SD", "unidad": "m"},
            "NF": {"nombre": "Nivel Freático", "unidad": "msnm"},
            "NI": {"nombre": "Nivel Incremental", "unidad": "m"},
            "NA": {"nombre": "Nivel Acumulado", "unidad": "m"},
            "PB": {"nombre": "Presión Barométrica", "unidad": "m"},
            "FP": {"nombre": "Frecuencia", "unidad": "m"},
            "TP": {"nombre": "Temperatura", "unidad": "m"},
            "VI": {"nombre": "Velocidad Incremental", "unidad": "m"},
            "AC": {"nombre": "Asentamiento en Cota", "unidad": "msnm"},
            "AI": {"nombre": "Asentamiento Incremental", "unidad": "m"},
            "AA": {"nombre": "Asentamiento Acumulado", "unidad": "m"},
            "UDI": {"nombre": "Umbral Desplazamiento", "unidad": "m"}
        }
        # Devolver el nombre y la unidad correspondiente al tipo
        info = tipo_a_info.get(tipo, {"nombre": "Tipo desconocido", "unidad": "N/A"})
        return info["nombre"], info["unidad"]
    